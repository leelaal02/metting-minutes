"""[4-공통] docx 저장 직전 후처리.

렌더러(무양식/템플릿)가 최종 저장하기 전에 공통으로 적용하는 마무리 작업.
표 행 나눔 허용과, 양식 글꼴을 렌더 결과 런에 입히는 작업을 한다.
"""
from docx.oxml.ns import qn


def _iter_all_tables(container):
    """문서/셀 안의 모든 표를 중첩표까지 재귀로 낸다.

    Document 본문과 `_Cell`은 모두 `.tables`를 가지므로 한 함수로 처리한다.
    """
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_tables(cell)


_FONT_TAGS = ("w:rFonts", "w:sz", "w:szCs")


def _iter_all_paragraphs(container):
    """문서/셀 안의 모든 문단을 중첩표까지 재귀로 낸다."""
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def inherit_mark_fonts(doc) -> int:
    """문단 부호에 적어 둔 양식 글꼴을, 글꼴이 없는 런에 입힌다.

    `apply_form_mapping`이 채운 문단에는 그 자리의 글꼴을 문단 부호(`w:pPr/w:rPr`)에
    남겨 둔다. 런 자체에도 넣지만, `{{r ... }}` 자리는 **렌더 시점에 docxtpl가 런을
    새로 만들어** 그 글꼴이 사라진다(굵은 소제목·빨간 "입력필요"). 그래서 저장 직전에
    표식을 보고 글꼴 없는 런에만 채워 넣는다 — 이미 글꼴이 있는 런은 건드리지 않는다.

    글꼴을 채운 런 수를 반환한다(테스트·로깅용).
    """
    import copy

    filled = 0
    for paragraph in _iter_all_paragraphs(doc):
        pPr = paragraph._p.find(qn("w:pPr"))
        mark = pPr.find(qn("w:rPr")) if pPr is not None else None
        if mark is None:
            continue
        wanted = [(tag, mark.find(qn(tag))) for tag in _FONT_TAGS]
        wanted = [(tag, el) for tag, el in wanted if el is not None]
        if not wanted:
            continue
        for run in paragraph.runs:
            rPr = run._r.get_or_add_rPr()
            changed = False
            for tag, element in wanted:
                if rPr.find(qn(tag)) is not None:
                    continue  # 런에 이미 지정된 글꼴이 우선
                if tag == "w:rFonts":
                    target = rPr.get_or_add_rFonts()
                    for name, value in element.attrib.items():
                        target.set(name, value)
                else:
                    rPr.append(copy.deepcopy(element))
                changed = True
            if changed:
                filled += 1
    return filled


def allow_rows_to_break(doc) -> int:
    """모든 표 행이 페이지 끝에서 나뉠 수 있게 한다(w:cantSplit 제거).

    많은 회의록 양식 템플릿이 행 나눔 금지(`w:cantSplit`)를 걸어 두어, 한 칸에
    긴 내용(회의 내용·논의 등)을 채우면 페이지 경계에서 행이 안 쪼개져 통째로
    다음 장으로 밀리거나 넘친다. Word 기본값(행 나눔 허용)에 맞춰 `cantSplit`를
    없애 내용이 페이지를 자연스럽게 넘어가게 한다 — 사용자가 매번 워드에서
    "페이지 끝에서 행 나눔 허용"을 켜지 않아도 되도록.

    제거한 행 수를 반환한다(테스트·로깅용).
    """
    removed = 0
    for table in _iter_all_tables(doc):
        for row in table.rows:
            trPr = row._tr.find(qn("w:trPr"))
            if trPr is None:
                continue  # 속성 자체가 없으면 이미 나눔 허용(기본값)
            for cant_split in trPr.findall(qn("w:cantSplit")):
                trPr.remove(cant_split)
                removed += 1
    return removed
