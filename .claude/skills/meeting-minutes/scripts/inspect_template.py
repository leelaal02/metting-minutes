"""[4-보강] 양식 구조 덤프: 토큰 없는 표 서식(.docx) → 구조 JSON.

자동 매핑 경로의 ①단계. 표·칸 라벨·좌표·빈칸·병합·음영·토큰 유무를 JSON으로
출력해 Claude가 10항목을 어느 칸에 넣을지 판단(mapping.json)하는 입력으로 쓴다.
스크립트는 순수 기계적(구조만 보고), 의미 판단은 하지 않는다.

**어디에 넣을지 판단하려면 "글자 유무"만으로는 부족하다.** 그래서 세 가지 배치
신호를 함께 낸다 — 이게 없으면 라벨 칸에 본문을 쓰거나 값 자리를 빈 채로 남긴다:
- `shaded`: 셀 배경색 유무(색 무관 — 회색·파랑·테마색 모두). 색칠된 칸은
  라벨/헤더이므로 값을 넣으면 안 된다.
- `numbered`: 문단의 자동 번호·글머리 유무. 빈 번호 문단은 "여백"이 아니라
  값을 적으라고 비워 둔 자리다(비워 두면 "1." "2."만 남는다).
- `blocks`: 표와 문단이 본문에 놓인 순서. 표 밖 문단이 어느 표 뒤인지 알아야
  "회의 내용" 라벨 다음 자리가 표 밖 문단이라는 걸 알 수 있다.

병합 셀은 원점(top-left) 1회만 출력하고, 병합 중복 셀은 제외한다.
"""
import json
import sys

from normalize_input import resolve_input_path


def _iter_container_texts(container):
    """컨테이너(본문·머리말/꼬리말·표 셀)의 문단 텍스트를 중첩표까지 재귀로 낸다.

    Document 본문·`_Header`/`_Footer`·`_Cell`은 모두 `.paragraphs`·`.tables`를
    공유하므로 한 함수로 처리한다. paragraph.text는 런을 이어붙이므로 한 문단 안의
    토큰은 쪼개지지 않는다. 셀 안의 중첩표까지 재귀해 어디에 있든 토큰을 놓치지 않는다.
    """
    for p in container.paragraphs:
        yield p.text
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_texts(cell)


def _iter_all_texts(doc):
    """문서 본문 + 모든 섹션의 머리말/꼬리말 텍스트를 재귀로 낸다(토큰 탐지용).

    토큰은 본문뿐 아니라 머리말/꼬리말·중첩표에도 들어갈 수 있다. 한 곳이라도
    놓치면 `has_tokens`가 False로 오판돼, 이미 토큰이 있는 양식이 자동 매핑 경로로
    잘못 흘러가 토큰이 이중 삽입된다. 그래서 탐지는 전 영역을 훑는다.
    """
    yield from _iter_container_texts(doc)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from _iter_container_texts(hf)


def _has_tokens(doc) -> bool:
    """문서 전체(본문·머리말/꼬리말·중첩표)에 docxtpl 토큰(`{{`/`{%`)이 있으면 True."""
    return any("{{" in t or "{%" in t for t in _iter_all_texts(doc))


# 배경색 없음으로 볼 fill 값(흰색·자동·미지정). 그 외는 색을 가리지 않고 음영으로 본다
# — 라벨 칸은 회색이 가장 흔할 뿐, 파랑·노랑·연두 등 어떤 색으로든 칠해질 수 있다.
_UNSHADED_FILLS = {None, "auto", "FFFFFF"}


def is_shaded(tc) -> bool:
    """표 셀(`w:tc`)에 배경 음영이 있으면 True — 라벨/헤더 칸 판별 신호.

    회의록 양식은 라벨 칸에 배경색을 주는 경우가 많다. 라벨이 같은 칸에 있는지
    (`"제목:" __`) 아니면 라벨 행 아래 흰 행이 값 자리인지는 글자 유무만으로
    구분되지 않으므로, 음영을 별도 신호로 낸다. `apply_form_mapping`도 이 함수로
    음영 칸 오배치를 막는다(단일 판정 기준).

    색상은 가리지 않는다. 세 가지 경로를 모두 본다:
    - `w:fill` — 직접 지정한 RGB(회색 F3F3F3, 파랑 4472C4 …)
    - `w:themeFill` — 테마 색으로 칠한 칸. 이때 `w:fill`은 비어 있어 fill만 보면 놓친다.
    - `w:val` — fill 없이 무늬(pct15·diagStripe 등)로 음영을 주는 경우.

    한계: 표 **스타일**(밴딩·첫 행 강조)에서 색이 오는 칸은 셀에 `w:shd`가 없어
    탐지되지 않는다. 그런 양식은 구조 JSON의 `text`·`blocks`로 라벨을 판단한다.
    """
    from docx.oxml.ns import qn

    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return False
    if shd.get(qn("w:fill")) not in _UNSHADED_FILLS:
        return True
    if shd.get(qn("w:themeFill")) is not None:
        return True
    return shd.get(qn("w:val")) not in (None, "clear", "nil")


def _is_numbered(paragraph) -> bool:
    """문단에 자동 번호/글머리(`w:numPr`)가 걸려 있으면 True.

    글자가 없어도 번호는 찍히므로, 이런 빈 문단을 안 채우면 결과물에 "1." "2."만
    남는다 — 간격용 빈 문단과 반드시 구분해야 한다.
    """
    from docx.oxml.ns import qn

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _inspect_blocks(doc) -> list:
    """본문의 표·문단이 놓인 순서를 `{"type", "index"}` 목록으로 낸다.

    `tables`/`paragraphs`는 각각 따로 번호가 매겨져 서로의 위치 관계를 알 수 없다.
    표 라벨 바로 다음 값 자리가 표 밖 문단인 양식이 있어 순서 정보가 필요하다.
    인덱스는 각각 `doc.tables`·`doc.paragraphs` 상의 위치와 일치한다.
    """
    blocks = []
    p_i = t_i = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            blocks.append({"type": "paragraph", "index": p_i})
            p_i += 1
        elif tag == "tbl":
            blocks.append({"type": "table", "index": t_i})
            t_i += 1
    return blocks


def _inspect_table(table, index: int) -> dict:
    """한 표의 셀 구조를 JSON 친화 dict로. 병합 원점만 출력.

    병합은 XML로 판정한다(lxml 프록시 id는 접근마다 달라 신뢰 못 함):
    - gridSpan > 1 → 가로 병합.
    - vMerge == 'restart' → 세로 병합 원점, 'continue' → 아래로 이어진 중복 셀(제외).
    셀의 논리 열은 tc를 왼쪽부터 훑으며 gridSpan만큼 누적해 계산한다.
    """
    from docx.table import _Cell

    n_rows = len(table.rows)
    n_cols = len(table.columns)

    cells = []
    for r, row in enumerate(table.rows):
        col = 0
        for tc in row._tr.tc_lst:
            span = tc.grid_span
            vmerge = tc.vMerge  # None | 'restart' | 'continue'
            if vmerge == "continue":
                col += span  # 세로 병합 중복 — 원점(위 restart 행)만 남긴다
                continue
            text = _Cell(tc, table).text.strip()
            cells.append({
                "row": r,
                "col": col,
                "text": text,
                "is_empty": text == "",
                "merged": span > 1 or vmerge == "restart",
                "shaded": is_shaded(tc),
            })
            col += span
    return {"index": index, "rows": n_rows, "cols": n_cols, "cells": cells}


def _inspect_paragraphs(doc) -> list:
    """본문 문단을 인덱스와 함께 덤프(문단 기반 양식 매핑용).

    인덱스는 doc.paragraphs 상의 위치이며 mapping.json의 `para` 주소와 일치한다.
    빈 문단도 인덱스 정확성을 위해 그대로 포함한다. 표 안의 문단은 제외된다.
    빈 문단이라도 `numbered`가 True면 간격용이 아니라 값을 적는 자리다.
    """
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        out.append({
            "index": i,
            "text": text,
            "is_empty": text == "",
            "numbered": _is_numbered(p),
        })
    return out


def inspect_template(template_path: str) -> dict:
    """.docx 양식을 열어 구조 JSON(dict)을 반환한다.

    표(`tables`)와 본문 문단(`paragraphs`)을 모두 덤프하므로, 표 기반·문단 기반
    양식 모두에 매핑을 만들 수 있다. `blocks`는 둘이 본문에 놓인 순서다.
    """
    from docx import Document

    resolved = resolve_input_path(template_path, must_exist=True)
    if resolved.suffix.lower() != ".docx":
        raise ValueError(
            f"양식은 .docx여야 합니다: {resolved.name}. "
            "hwp/pdf 양식이면 한글/워드에서 .docx로 저장해 다시 주세요."
        )
    doc = Document(str(resolved))
    return {
        "has_tokens": _has_tokens(doc),
        "blocks": _inspect_blocks(doc),
        "tables": [_inspect_table(t, i) for i, t in enumerate(doc.tables)],
        "paragraphs": _inspect_paragraphs(doc),
    }


def main() -> None:
    path = sys.argv[1]
    result = inspect_template(path)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
