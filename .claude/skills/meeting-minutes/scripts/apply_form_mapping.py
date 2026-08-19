"""[4-보강] 매핑 적용: 양식 복사본에 검증된 docxtpl 토큰 블록을 삽입.

자동 매핑 경로의 ③단계. Claude가 만든 mapping.json대로 표의 지정 칸에
**검증된 토큰 텍스트만** 넣는다(Claude는 토큰 문법을 짜지 않는다 → 문법 오류 차단).
삽입 후에는 기존 render_docx_template.py가 무수정으로 값을 채운다.

원본 양식은 절대 수정하지 않는다 — 호출자가 준 복사본 경로(out_path)에만 저장.
"""
import copy
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from inspect_template import is_shaded
from normalize_input import resolve_input_path

# --- 토큰 블록 라이브러리 (설계 §4) ------------------------------------------
# 서식(빨강 "입력필요"·소제목 굵게)이 필요한 값은 RichText 슬롯이라 반드시
# `{{r ... }}`(r 접두사)로 넣는다. 일반 `{{ }}`에 RichText를 주면 run이 중첩되어
# 빈칸으로 깨진다(docxtpl 제약). 대응 컨텍스트 키는 render_docx_template.build_context
# 의 `*_rt`/`discussion_rt`/`action_items_rt`. 항상-존재하는 평문(points·task 등)은
# 일반 `{{ }}`로 둔다.
_SCALAR_TOKEN = {
    "title": "{{r title_rt }}",
    "date": "{{r date_rt }}",
    "attendees": "{{r attendees_rt }}",
    "purpose": "{{r purpose_rt }}",
    "next_meeting": "{{r next_meeting_rt }}",
}
_TODO_TOKEN = "{{r todo }}"  # build_context의 RichText `todo`("입력필요" 빨강·굵게)

# 한 칸에 여러 field가 들어갈 때 구분용 섹션 라벨(block 복수 field에서만 붙임).
_SECTION_LABEL = {
    "title": "[제목]",
    "date": "[일시]",
    "attendees": "[참석자]",
    "purpose": "[회의 목적]",
    "next_meeting": "[다음 회의]",
    "discussion": "[논의 내용]",
    "decisions": "[결정 사항]",
    "open_issues": "[미결 사항]",
    "action_items": "[실행 항목]",
    "notes": "[기타·특이사항]",
}

# 허용되는(Field) 목록을 모아 놓은 상수
ALLOWED_FIELDS = set(_SECTION_LABEL)  # 고정 10항목 어휘

# --- 계층 기호 사다리 --------------------------------------------------------
# 양식이 이미 쓰는 기호를 우리 내용이 또 쓰면 두 계층이 같은 기호가 돼 구분이
# 안 된다("1. [논의 내용]" 아래 "1. 주제"). 그래서 기호를 고정하지 않고,
# **그 자리에서 양식이 쓰는 기호를 비껴** 사다리에서 골라 쓴다.
_TOPIC_LADDER = ["num", "□", "◇", "▪"]  # "num" = 렌더러의 "N. 주제" 넘버링 # 주제
_ITEM_LADDER = ["-", "·", "‣", "◦"] # 항목, 내용
_ITEM_MARKS = set(_ITEM_LADDER)  # _line_level이 라인만 보고 계층을 판정하도록
_DEFAULT_MARKERS = {"topic": _TOPIC_LADDER[0], "item": _ITEM_LADDER[0]} # 기본 사용 기호: 주제 1.  항목 -


# 행 반복 표(row_repeats)용: 리스트형 field → 반복 대상 컨텍스트 키·루프 변수·
# 하위필드별 컬럼 토큰. task는 항상-존재 평문({{ }}), owner/due는 RichText 슬롯({{r }}).
_ROW_REPEAT = { # 반복 가능한 표 설정
    "action_items": {
        "iter": "action_items_rt", # 템플릿에 전달할 반복 데이터 이름
        "var": "a", # 반복문에서 사용할 변수 이름
        "col_tokens": { # 각 열(Column)에 넣을 토큰
            "task": "{{ a.task }}", # 해야 할 작업(업무 내용)
            "owner": "{{r a.owner_rt }}", # 담당자
            "due": "{{r a.due_rt }}",# 완료 예정일(Due Date)
        },
    },
}


def _list_block(field: str, markers: dict) -> list:
    """목록형 field의 block 본문(문단 라인 목록). {%p%}는 문단 단위 반복 태그."""
    item = markers["item"] # 항목
    topic = markers["topic"] # 주제
    # 넘버링: 렌더러가 만든 "N. 주제",  기호: 기호 + 주제(굵게, 번호 없음).
    topic_line = (
        "{{r d.topic_rt }}" if topic == "num" else f"{topic} {{{{r d.topic_plain_rt }}}}"
    )
    blocks = { # 필드별 출력 블록을 저장한 딕셔너리
        "discussion": [ # 논의 내용
            "{%p for d in discussion_rt %}",
            topic_line,
            "{%p for p in d.points %}",
            f" {item} {{{{ p }}}}",
            "{%p endfor %}",
            "{%p endfor %}",
        ],
        "decisions": [ # 결정 사항
            "{%p for x in decisions %}",
            f" {item} {{{{ x }}}}",
            "{%p endfor %}",
        ],
        "open_issues": [ # 미결 사항
            "{%p for x in open_issues %}",
            f" {item} {{{{ x }}}}",
            "{%p endfor %}",
        ],
        "action_items": [ # 실행 항목
            "{%p for a in action_items_rt %}",
            f" {item} {{{{ a.task }}}} (담당: {{{{r a.owner_rt }}}} / 기한: {{{{r a.due_rt }}}})",
            "{%p endfor %}",
        ],
        "notes": [ # 비고
            "{%p for n in notes %}",
            f" {item} {{{{ n }}}}",
            "{%p endfor %}",
        ],
    }
    return blocks[field]

# fields에 들어있는 필드 이름이 모두 올바른지 검사(검증)하는 함수
def _validate_fields(fields) -> None:
    if not fields: # if 조건문 : 조건이 참일때만 실행
        raise ValueError("fill의 'fields'가 비어 있습니다.")
    for f in fields: # for 반복문 : 여러개 요소를 하나씩 반복 처리
        if f not in ALLOWED_FIELDS:
            raise ValueError( # raise : 예외발생 오류를 발생시켜 실행 중단
                f"알 수 없는 field '{f}'. 허용 어휘: "
                + ", ".join(sorted(ALLOWED_FIELDS))
            )


def inline_tokens(fields) -> str:
    """inline mode: 스칼라 field의 값 토큰을 공백으로 이어 한 줄로."""
    _validate_fields(fields)
    parts = []
    for f in fields:
        if f not in _SCALAR_TOKEN:
            raise ValueError(
                f"'{f}'는 목록형이라 inline 모드로 넣을 수 없습니다. block 모드를 쓰세요."
            )
        parts.append(_SCALAR_TOKEN[f])
    return " ".join(parts)


def block_lines(fields, markers: dict = None) -> list:
    """block mode: field 순서대로 토큰 문단 라인 목록을 만든다.

    field가 하나면 섹션 라벨 생략, 둘 이상이면 각 field에 섹션 라벨 줄을 붙인다.
    `markers`로 계층 기호를 바꿀 수 있다(양식과 겹치지 않게 자동 선택된 값).
    """
    _validate_fields(fields)
    markers = markers or _DEFAULT_MARKERS
    single = len(fields) == 1
    lines = []
    for f in fields:
        if f in _SCALAR_TOKEN:  # 스칼라·title·date는 값 토큰 한 줄
            token = _SCALAR_TOKEN[f]
            lines.append(token if single else f"{_SECTION_LABEL[f]} {token}")
        else:  # 목록형: 반복 블록
            if not single:
                lines.append(_SECTION_LABEL[f])
            lines.extend(_list_block(f, markers))
    return lines


# --- 글꼴 읽기·입히기 --------------------------------------------------------
_FONT_ATTRS = ("ascii", "hAnsi", "eastAsia", "cs")


def read_font(paragraph):
    """문단의 런에서 글꼴 이름·크기를 읽는다(없으면 None).

    양식은 글꼴을 스타일이 아니라 **런에 직접** 지정하는 경우가 많다(예:
    "ㅇ (목적)"이 HCI Poppy 14pt). 우리가 새로 만든 런은 서식이 비어 있어
    스타일 기본값으로 렌더되므로, 같은 자리에 글꼴이 섞인다. 그래서 그 자리의
    글꼴을 읽어 우리가 채우는 런에 그대로 입힌다.
    """
    # Paragraph : 문단,  Run : 같은 서식의 텍스트
    sources = [run._r.find(qn("w:rPr")) for run in paragraph.runs] # sources : 글꼴 정보를 담고 있는 rPr를 모아 놓은 리스트 
    pPr = paragraph._p.find(qn("w:pPr")) # pPr : 문단(Paragraph)의 속성 
    if pPr is not None:  # 런을 지운 뒤에도 남는 문단 부호 서식(예시 문구 제거 대비)
        sources.append(pPr.find(qn("w:rPr"))) # sources라는 리스트의 append() 함수를 호출
    for rPr in sources:
        if rPr is None:
            continue
        names = {}
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr in _FONT_ATTRS:
                value = rFonts.get(qn(f"w:{attr}"))
                if value:
                    names[attr] = value
        sz = rPr.find(qn("w:sz"))
        size = int(sz.get(qn("w:val"))) / 2 if sz is not None else None
        if names or size:
            return {"names": names, "size": size}
    return None


def _write_font(rPr, font) -> None:
    """`w:rPr`에 글꼴 이름·크기를 쓴다(python-docx API로 순서 안전하게)."""
    if font["names"]:
        rFonts = rPr.get_or_add_rFonts()
        for attr, value in font["names"].items():
            rFonts.set(qn(f"w:{attr}"), value)
    if font["size"]:
        rPr.sz_val = Pt(font["size"])


def apply_font(run, font) -> None:
    """런에 글꼴을 입힌다(굵기·색 등 다른 서식은 건드리지 않는다)."""
    if not font:
        return
    _write_font(run._r.get_or_add_rPr(), font)


def mark_font(paragraph, font) -> None:
    """문단 부호(`w:pPr/w:rPr`)에 글꼴을 적어 둔다 — 렌더 뒤 후처리용 표식.

    `{{r ... }}`(RichText) 자리는 렌더 시점에 docxtpl가 런을 새로 만들기 때문에
    지금 런에 글꼴을 넣어도 사라진다. 그래서 "이 문단은 이 글꼴" 정보를 문단
    부호에 남겨 두고, `docx_postprocess.inherit_mark_fonts`가 렌더 후에 입힌다.
    문단 부호의 rPr은 글자 렌더에 영향을 주지 않으므로 안전한 보관 장소다.
    """
    if not font:
        return
    pPr = paragraph._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        # rPr은 pPr 안에서 sectPr·pPrChange 바로 앞이 제자리다(스키마 순서).
        pPr.insert_element_before(rPr, "w:sectPr", "w:pPrChange")
    _write_font(rPr, font)


# --- 문단 조작 ---------------------------------------------------------------
def _clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _append_token(paragraph, token: str) -> None:
    """paragraph 끝에 임의의 토큰/텍스트를 덧붙인다(라벨 텍스트 보존).

    라벨이 쓰는 글꼴을 그대로 물려받는다 — 안 그러면 "제 목 :"은 양식 글꼴,
    그 뒤 값만 스타일 기본 글꼴로 렌더돼 한 줄 안에서 글꼴이 갈린다.
    """
    font = read_font(paragraph)
    sep = " " if paragraph.text.strip() else ""  # 라벨이 있으면 한 칸 띄움
    run = paragraph.add_run(sep + token)
    apply_font(run, font)
    mark_font(paragraph, font)


# 양식이 인원수를 적으라고 비워 둔 자리("총  명 참석"). 숫자만 빠져 있어 값 토큰을
# 뒤에 붙이는 것으로는 못 채운다 — 빈칸 자체에 수를 끼워 넣어야 문장이 완성된다.
_COUNT_GAP = re.compile(r"총(\s+)명")


def _fill_attendee_count_gap(paragraph) -> bool:
    """"총 __ 명" 빈칸에 참석인원 수 토큰을 끼워 넣는다(참석자 inline 자리에서만).

    수는 세지 않고 `attendees` 길이에서 나오므로(build_context) 명단과 어긋날 수
    없다. 공백이 없는 "총명"은 빈칸이 아니므로 건드리지 않는다.

    **문단 전체 텍스트에서 찾는다** — 워드는 한 문장을 여러 런으로 쪼개 두는 일이
    잦아("총"/"  "/"명" 3개 런) 런 하나만 보면 빈칸을 놓친다. 찾은 구간에 걸친
    런들만 다시 쓰고 나머지 런은 건드리지 않아 양식 서식이 그대로 남는다.
    """
    m = _COUNT_GAP.search(paragraph.text)
    if not m:
        return False
    start, end = m.span()
    filled = False
    pos = 0
    for run in paragraph.runs:
        r_start, r_end = pos, pos + len(run.text)
        pos = r_end
        if r_end <= start or r_start >= end:  # 매치 밖 런은 그대로 둔다
            continue
        head = run.text[: max(0, start - r_start)]
        tail = run.text[end - r_start :] if r_end > end else ""
        # 토큰은 첫 런에만 넣고(그 런의 글꼴을 물려받는다), 나머지 걸친 런은 비운다.
        run.text = head + ("" if filled else "총 {{ attendee_count }}명") + tail
        filled = True
    return filled


def _copy_paragraph_format(source, target) -> None:
    """원본 문단의 서식(`w:pPr`)을 복제하되 번호 매김(`w:numPr`)은 뺀다.

    새로 만든 문단은 pPr이 없어 문서 기본 서식(Normal)이 된다 — 양식의 글꼴·
    정렬·줄간격이 채운 부분에서만 튀는 원인. 반대로 numPr까지 복사하면 삽입한
    **모든 줄에 번호가 붙으므로** 제외한다(번호 자리에 목록을 넣는 양식 대응).
    """
    src_pPr = source._p.find(qn("w:pPr"))
    if src_pPr is None:
        return
    new_pPr = copy.deepcopy(src_pPr)
    for num_pr in new_pPr.findall(qn("w:numPr")):
        new_pPr.remove(num_pr)
    old_pPr = target._p.find(qn("w:pPr"))
    if old_pPr is not None:
        target._p.remove(old_pPr)
    target._p.insert(0, new_pPr)


def _insert_paragraph_after(paragraph, text: str, style_from=None):
    """본문에서 paragraph 바로 뒤에 새 문단을 만들어 반환한다.

    python-docx에 공개 API가 없어 XML(addnext)로 삽입한다. 표 셀이 아니라
    문단 기반 양식의 block 채움에서 여러 문단을 순서대로 끼워 넣을 때 쓴다.
    `style_from`을 주면 그 문단의 서식을 물려받는다(양식 글꼴 유지).
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_from is not None:
        _copy_paragraph_format(style_from, new_para)
    if text:
        new_para.add_run(text)
    return new_para


# --- 자리 판정(라벨·예시 문구·계층 기호) -------------------------------------
# 글머리 기호만 있는 자리("ㅇ", "-", "1.")는 라벨이 아니라 빈 값 자리다.
_BULLET_CHARS = "ㅇ○●◦·•*-–—□■◇◆※∙"
_NUMBER_PREFIX = re.compile(r"^\(?\d+\s*[.)]\s*")

# 양식이 미리 넣어 둔 예시 문구 패턴. 라벨이 아니라 지우고 채워야 할 자리다.
_ASK_PATTERN = re.compile(r"(작성|입력|기재|기입)\s*(해\s*주)?\s*(하)?세요")
_FILLER_RUN = re.compile(r"[Oo0○]{2,}")


def is_placeholder(text: str) -> bool:
    """양식의 예시 문구면 True — "내용을 작성하세요.", "OO팀 OOO", "2025년 00월 00일".

    라벨("회의일시")은 자리를 알려주므로 보존해야 하지만, 예시 문구는 **지우고
    그 자리에 값을 넣어야** 한다. 남겨 두면 회의록에 "내용을 작성하세요."가 그대로
    인쇄된다. 판정은 두 가지 — 작성 요청 문구, 그리고 O·0을 반복한 빈칸 표기.
    """
    stripped = text.strip()
    if not stripped:
        return False
    if _ASK_PATTERN.search(stripped):
        return True
    return len(_FILLER_RUN.findall(stripped)) >= 2


def _has_meaningful_label(text: str) -> bool:
    """그 자리에 "무엇을 적는 칸인지 알려주는 라벨"이 있으면 True.

    "결정 사항:"은 라벨이지만 "ㅇ"·"-"·"1."은 글머리 기호일 뿐이다. 글머리만
    있는 자리에 목록을 넣으면 **제목 없이 내용만 남아** 무슨 항목인지 알 수 없으므로,
    이 판정으로 제목(섹션 라벨)을 자동으로 붙일지 정한다.
    """
    stripped = _NUMBER_PREFIX.sub("", text.strip().strip(_BULLET_CHARS + " "))
    return bool(stripped.strip(_BULLET_CHARS + " "))


def _slot_markers(paragraph) -> set:
    """그 자리(문단)가 이미 쓰고 있는 계층 기호를 모은다.

    양식이 "ㅇ"·"-"·자동번호를 쓰고 있으면 우리 내용이 같은 기호를 쓰면 안 된다.
    자동번호(`w:numPr`)와 "1." 같은 글자 번호는 둘 다 `"num"`으로 본다.
    """
    used = set()
    text = paragraph.text.strip()
    if text:
        head = text[0]
        if head in _BULLET_CHARS:
            used.add(head)
        if _NUMBER_PREFIX.match(text):
            used.add("num")
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        used.add("num")
    return used


def _form_number_survives(target, first_line: str, label) -> bool:
    """양식의 자동 번호가 **결과물에 실제로 남는지** 판정한다.

    남지 않는 번호를 피하면 쓸 수 있는 넘버링("1. 주제")을 공짜로 버리고 사다리
    아래 기호(`□`)로 내려간다 — 겹치지도 않는 충돌을 피하는 손해다. 세 경우로 갈린다.

    - 제목을 넣는 자리: `_put_section_label`이 번호를 지우므로 남지 않는다.
    - 라벨을 보존하는 자리("결정 사항:"): 라벨과 함께 번호도 그대로 남는다.
    - 빈 자리에 첫 줄을 덮어쓰는 자리: 그 줄이 `{%p … %}` 태그면 렌더 시 **문단째
      삭제**되므로 번호도 사라진다. 태그가 아닌 값 토큰이면 번호가 남는다.
    """
    if label:
        return False
    if target.text.strip():
        return True
    return not first_line.lstrip().startswith("{%")


def _choose_markers(used: set) -> dict:
    """양식이 쓰는 기호를 비껴 우리 계층 기호를 고른다.

    사다리를 앞에서부터 훑어 겹치지 않는 첫 기호를 쓴다. 양식이 사다리를 다
    쓰고 있으면(=우리가 쓸 계층이 양식보다 깊으면) 마지막 기호로 떨어지되,
    들여쓰기가 한 단 더 들어가므로 계층은 여전히 구분된다.
    """
    return {
        "topic": next((m for m in _TOPIC_LADDER if m not in used), _TOPIC_LADDER[-1]),
        "item": next((m for m in _ITEM_LADDER if m not in used), _ITEM_LADDER[-1]),
    }


# --- 블록 라인 서식 ----------------------------------------------------------
def _line_level(line: str):
    """블록 라인의 계층을 판정한다 — 0=섹션 제목, 1=소제목·값, 2=하위 항목, None=태그.

    **계층마다 기호가 달라야 한다** — 같은 기호가 두 계층에 나오면(예: 섹션 제목이
    양식 번호를 물려받아 "1. [논의 내용]"이 되고 그 아래 주제도 "1. …") 어디까지가
    상위인지 읽히지 않는다. 그래서 `[섹션 제목]` / `1. 소제목` / `- 항목` 세 기호를
    쓰고 들여쓰기도 단계별로 벌린다.

    `{%p %}` 태그 문단은 렌더 시 삭제되므로 서식을 줄 필요가 없다.
    """
    stripped = line.strip()
    if stripped.startswith("{%"):
        return None
    if stripped.startswith("["):  # "[논의 내용]" 등 섹션 제목
        return 0
    if stripped[:1] in _ITEM_MARKS:  # "-"·"·"·"‣"·"◦" 등 하위 항목 기호
        return 2
    return 1  # 소제목("1. 주제"·"□ 주제")·스칼라 값


def _style_block_line(paragraph, line: str, depth: int = 0, font=None) -> None:
    """블록 라인에 계층 들여쓰기·여백을 준다 — 가독성의 핵심.

    들여쓰기가 없으면 소제목과 하위 항목이 같은 왼쪽 선에서 시작해 계층이
    보이지 않고, **내어쓰기(hanging)가 없으면 줄바꿈된 둘째 줄이 글머리 아래가
    아니라 맨 왼쪽에 붙어** 항목 경계가 사라진다. 한 항목이 두세 줄로 넘어가는
    회의록에서 특히 크게 티가 난다.
    """
    level = _line_level(line)
    if level is None:
        return
    for run in paragraph.runs:  # 양식 글꼴 물려받기
        apply_font(run, font)
    mark_font(paragraph, font)  # RichText 자리는 렌더 후 후처리가 입힌다
    # 양식이 이미 한 계층을 쓰고 있으면(글머리 칸) 그 아래로 한 단 더 들어간다.
    offset = Cm(0.4) * depth
    pf = paragraph.paragraph_format
    pf.space_after = Pt(0)
    if level == 2:  # "- 항목": 들여쓰고 내어쓰기로 둘째 줄을 글머리에 맞춤
        pf.left_indent = Cm(1.1) + offset
        pf.first_line_indent = Cm(-0.35)
        pf.space_before = Pt(0)
    elif level == 1:  # "1. 주제"·값: 섹션 제목보다 한 단 들여씀
        pf.left_indent = Cm(0.4) + offset
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
    else:  # "[논의 내용]" 섹션 제목: 맨 왼쪽·굵게·위로 크게 띄움
        pf.left_indent = Cm(0) + offset
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(10)
        for run in paragraph.runs:
            run.bold = True


def _put_section_label(paragraph, label: str, *, append: bool, depth=0, font=None):
    """섹션 제목을 문단에 넣는다 — 자동 번호 제거 + 굵게 + 위 여백.

    번호를 지우는 이유: 제목을 양식의 번호 자리("1." "2.")에 넣으면
    "1. [논의 내용]"이 되어 바로 아래 주제 번호("1. 컨설팅…")와 **같은 기호가 두
    계층에 겹친다.** `append=True`면 글머리("ㅇ") 뒤에 이어 붙이고, False면 빈
    문단을 제목으로 쓴다.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        for num_pr in pPr.findall(qn("w:numPr")):
            pPr.remove(num_pr)
    if append:
        _append_token(paragraph, label)
    else:
        _clear_runs(paragraph)
        paragraph.add_run(label)
    _style_block_line(paragraph, label, depth, font)


# --- block 삽입 --------------------------------------------------------------
def _apply_block(target, fields, add_line, *, labeled: bool = False) -> None:
    """block 토큰을 target 문단(표 셀 첫 문단 | 본문 문단) 자리에 삽입한다.

    라벨이 있는 자리("결정 사항:"·"ㅇ 논의내용")는 지우지 않고 그 뒤에 블록을 새
    문단으로 넣는다. 빈 자리면 첫 라인을 그 문단에 넣어 서식을 유지한다.
    라벨이 없는 자리("ㅇ"·빈 번호 문단)에는 섹션 제목을 자동으로 붙인다.

    `labeled=True`는 **그 자리 밖에** 이미 라벨이 있다는 뜻이다(옆 칸 "논의 내용" 등).
    자리 안의 글자만 보면 라벨이 없어 보여 제목이 또 붙는다.

    `add_line(line)`은 다음 라인을 놓을 위치를 아는 콜백(표 셀 끝 / 직전 문단 뒤)
    으로, 만든 문단을 돌려준다 — 표·본문의 차이는 이 콜백에만 있다.
    """
    _validate_fields(fields)  # block_lines보다 먼저 — 어휘 오류를 삽입 전에 낸다
    font = read_font(target)  # 그 자리가 쓰는 글꼴(런에 직접 지정된 값)
    # field가 둘 이상이면 block_lines가 항목마다 라벨을 붙이므로 제목은 생략.
    label = (
        _SECTION_LABEL[fields[0]]
        if len(fields) == 1 and not labeled and not _has_meaningful_label(target.text)
        else None
    )
    # 결과물에 **실제로 남을** 양식 기호만 피하면 된다. 제목 자리의 자동 번호나
    # 태그 문단의 번호는 렌더 전에 사라지므로, 여기서까지 피하면 넘버링을 버린다.
    used = _slot_markers(target)
    if not _form_number_survives(target, block_lines(fields)[0], label):
        used.discard("num")
    lines = block_lines(fields, _choose_markers(used))
    depth = 1 if used else 0  # 양식이 한 계층을 쓰면 그 아래로 들여쓴다
    if label:  # 제목이 필요한 자리 → 그 문단은 제목, 블록은 전부 그 뒤에
        _put_section_label(
            target, label, append=bool(target.text.strip()), depth=depth, font=font
        )
        rest = lines
    elif target.text.strip():  # 라벨 보존: 기존 텍스트 유지, 블록은 전부 뒤에
        rest = lines
    else:  # 빈 자리: 첫 라인은 그 문단에(서식 유지), 나머지는 뒤에
        _clear_runs(target)
        target.add_run(lines[0])
        _style_block_line(target, lines[0], depth, font)
        rest = lines[1:]
    for line in rest:
        _style_block_line(add_line(line), line, depth, font)


def _apply_block_cell(cell, fields, *, labeled: bool = False) -> None:
    """표 셀 채움 — 이어지는 라인은 셀 끝에 문단을 더한다."""
    p0 = cell.paragraphs[0]

    def add_line(line):
        added = cell.add_paragraph(line)
        _copy_paragraph_format(p0, added)  # 셀 서식 상속(스타일·정렬)
        return added

    _apply_block(p0, fields, add_line, labeled=labeled)


def _apply_block_paragraph(paragraph, fields, *, labeled: bool = False) -> None:
    """본문 문단 채움 — 이어지는 라인은 직전 문단 뒤에 XML로 끼워 넣는다."""
    prev = [paragraph]

    def add_line(line):
        # 서식은 항상 원본 문단에서 물려받는다(직전 삽입분이 아니라).
        prev[0] = _insert_paragraph_after(prev[0], line, style_from=paragraph)
        return prev[0]

    _apply_block(paragraph, fields, add_line, labeled=labeled)


# --- 매핑 항목 적용 ----------------------------------------------------------
def _require(entry: dict, key: str, where: str):
    """매핑 항목에서 필수 키를 꺼낸다. 없으면 원시 KeyError 대신 안내 메시지."""
    if key not in entry:
        raise ValueError(
            f"매핑 항목({where})에 필수 키 '{key}'가 없습니다. "
            "각 fills 항목은 row·col·mode·fields를, paragraphs 항목은 para·mode·fields를 가져야 합니다."
        )
    return entry[key]


def _at(seq, idx, what: str, where: str):
    """범위를 검사하고 seq[idx]를 돌려준다 — 원시 IndexError 대신 안내 메시지."""
    if idx < 0 or idx >= len(seq):
        raise IndexError(
            f"{what} {idx}이(가) 범위를 벗어났습니다({where}는 {len(seq)}개)."
        )
    return seq[idx]


def _check_not_shaded(cell, entry: dict, where: str) -> None:
    """배경색이 칠해진 칸에 값을 넣으려 하면 막는다(색 무관 — 회색뿐 아니라 어떤 색이든).

    색칠된 칸은 거의 항상 라벨/헤더이고, 값은 인접한 색 없는 칸이나 다음 행·문단에
    들어간다. block/inline은 라벨을 지우지 않고 뒤에 이어 붙이므로 색칠된 칸에
    매핑해도 오류 없이 통과해 **라벨 칸 안에 본문이 박히는** 결과가 나온다.
    이 조용한 오배치를 실행 시점에 시끄럽게 만든다.

    매핑 항목에 `"allow_shaded": true`가 있으면 통과시키되, 이 우회는
    **사용자가 그 칸에 넣으라고 명시적으로 요청했을 때만** 쓴다. 넣을 자리를
    못 찾았다는 이유로 붙이는 용도가 아니다(그 경우 다른 빈 자리로 옮긴다).
    """
    if entry.get("allow_shaded"):
        return
    if is_shaded(cell._tc):
        raise ValueError(
            f"{where}: 배경색이 칠해진 칸에 값을 넣으려 합니다. 색칠된 칸은 "
            "라벨/헤더이므로 값은 같은 행의 색 없는 칸이나 바로 다음 빈 행·문단에 "
            "매핑하세요(구조 JSON의 shaded=false 자리). 정말 이 칸에 넣어야 하면 "
            '해당 매핑 항목에 "allow_shaded": true를 추가하세요.'
        )


def is_duplicate_hint(table, row_index: int, cell) -> bool:
    """같은 행에 똑같은 글자가 또 있으면 그 칸은 라벨을 복사해 둔 힌트다.

    "부서 | 부서", "작성자 | 작성자"처럼 라벨 칸과 값 칸에 같은 말을 넣어 둔 양식이
    있다. 그대로 두면 "부서 입력필요"처럼 겹쳐 보이므로 예시와 똑같이 지운다.
    병합 칸은 `row.cells`가 같은 칸을 여러 번 돌려주므로 XML 요소로 자기 자신을 건다.
    """
    text = cell.text.strip()
    if not text:
        return False
    return any(
        other._tc is not cell._tc and other.text.strip() == text
        for other in table.rows[row_index].cells
    )


def has_row_label(table, row_index: int, cell) -> bool:
    """같은 행의 다른 칸이 이 칸의 라벨 역할을 하고 있으면 True.

    한국 회의록 양식의 표준 배치는 `논의 내용 | (빈칸)`처럼 라벨 칸과 값 칸을
    나란히 두는 것이다. 값 칸 안의 글자만 보면 라벨이 없어 보여 섹션 제목이
    자동으로 붙고, 결과가 `논의 내용 | [논의 내용] …`으로 같은 말이 두 번 나온다.
    그래서 자리 안뿐 아니라 **같은 행**도 함께 보고 제목을 붙일지 정한다.

    라벨이 값 칸 위쪽 행에 있는 양식(라벨 행 + 아래 빈 행)은 행이 달라 여기서
    잡히지 않는다 — 그런 자리는 매핑에서 `"auto_label": false`로 제목을 끈다.
    병합 칸은 `row.cells`가 같은 칸을 여러 번 돌려주므로 XML 요소로 자기 자신을 건다.
    """
    return any(
        other._tc is not cell._tc and _has_meaningful_label(other.text)
        for other in table.rows[row_index].cells
    )


def _resolve_labeled(entry: dict, detected: bool) -> bool:
    """섹션 제목을 붙일지 정한다 — 매핑의 `auto_label`이 자동 판정보다 우선.

    자동 판정(같은 행 라벨 유무)으로 대부분 맞지만, 라벨이 위쪽 행에 있거나
    옆 칸 라벨이 이 항목과 무관한 양식도 있다. 구조 JSON을 보는 쪽이 사람이므로
    `"auto_label": false`(제목 끄기)·`true`(제목 강제)로 뒤집을 수 있게 둔다.
    """
    override = entry.get("auto_label")
    return detected if override is None else not override


def _clear_placeholder(paragraph, *, text=None, extras=(), force: bool = False) -> None:
    """예시 문구가 든 자리를 비운다 — 글꼴은 문단 부호에 남겨 값이 같은 서식으로.

    판정 대상은 `text`(표 셀이면 칸 전체 글자), 없으면 문단 글자다. 표 셀은 예시가
    여러 줄인 경우가 있으므로("1. 내용을 작성하세요." ×4) 첫 문단만 남기고 `extras`
    문단은 통째로 제거한다. `force`면 문구 판정 없이 비운다.
    """
    if not force and not is_placeholder(paragraph.text if text is None else text):
        return
    font = read_font(paragraph)
    for extra in list(extras):
        extra._p.getparent().remove(extra._p)
    _clear_runs(paragraph)
    mark_font(paragraph, font)


def _apply_mode(entry: dict, paragraph, where: str, block_apply) -> None:
    """mode에 따라 토큰을 넣는다(inline|block|todo|literal) — 표·문단 공통 분기.

    block만 표(셀 끝에 문단 추가)와 본문(문단 뒤에 삽입)의 방식이 달라 콜백으로 받는다.
    """
    mode = _require(entry, "mode", where)
    if mode == "inline":
        fields = _require(entry, "fields", where)
        if "attendees" in fields:  # "총 __ 명" 빈칸이 있으면 인원수부터 채운다
            _fill_attendee_count_gap(paragraph)
        _append_token(paragraph, inline_tokens(fields))
    elif mode == "block":
        block_apply(_require(entry, "fields", where))
    elif mode == "todo":  # 라벨만 있고 데이터가 없는 자리 → 빨간 "입력필요"
        _append_token(paragraph, _TODO_TOKEN)
    elif mode == "literal":  # 원문에서 찾은 값 등 지정 문자열을 평문으로
        _append_token(paragraph, _require(entry, "text", where))
    else:
        raise ValueError(
            f"알 수 없는 mode '{mode}' (inline|block|todo|literal 중 하나)."
        )


def _apply_table_fills(doc, mapping) -> None:
    """표 셀 채움(`fills`).

    표는 항목별 `table` 키로 고르고, 없으면 최상위 `table`(기본 0)을 쓴다.
    한 양식에 표가 여러 개여도 한 번의 실행으로 채울 수 있다.
    """
    default_ti = mapping.get("table", 0)
    for fill in mapping.get("fills", []):
        ti = fill.get("table", default_ti)
        table = _at(doc.tables, ti, "표 인덱스", "문서의 표")
        r = _require(fill, "row", "fills")
        c = _require(fill, "col", "fills")
        _require(fill, "mode", "fills")  # 문서를 건드리기 전에 키 누락을 알린다
        row_cells = _at(table.rows, r, "행", f"표 {ti}의 행").cells
        cell = _at(row_cells, c, "열", f"행 {r}의 열")
        _check_not_shaded(cell, fill, f"fills[표{ti} {r},{c}]")
        # 예시 문구·라벨 복사 힌트는 지우고 그 자리에 값을 넣는다.
        _clear_placeholder(
            cell.paragraphs[0],
            text=cell.text,
            extras=cell.paragraphs[1:],
            force=is_duplicate_hint(table, r, cell),
        )
        # 옆 칸이 라벨인 배치(`논의 내용 | (빈칸)`)에서는 섹션 제목을 붙이지 않는다.
        labeled = _resolve_labeled(fill, has_row_label(table, r, cell))
        _apply_mode(
            fill,
            cell.paragraphs[0],
            "fills",
            lambda f: _apply_block_cell(cell, f, labeled=labeled),
        )


def _apply_paragraph_fills(doc, mapping) -> None:
    """문단 채움(`paragraphs`). 인덱스를 먼저 스냅샷해 삽입에 따른 인덱스 밀림을 피한다."""
    paras = doc.paragraphs
    # block 삽입이 뒤 인덱스를 밀기 전에 대상 문단 객체를 먼저 확보한다.
    targets = [
        (_at(paras, _require(pf, "para", "paragraphs"), "문단 인덱스", "본문 문단"), pf)
        for pf in mapping.get("paragraphs", [])
    ]
    for paragraph, pf in targets:
        _clear_placeholder(paragraph)  # 예시 문구는 지우고 채운다
        # 본문 문단은 옆 칸이 없어 자동 판정 대상이 아니다(자리 안의 라벨만 본다).
        # 앞 문단이 라벨인 양식은 매핑에서 `"auto_label": false`로 끈다.
        labeled = _resolve_labeled(pf, False)
        _apply_mode(
            pf,
            paragraph,
            "paragraphs",
            lambda f, p=paragraph, lb=labeled: _apply_block_paragraph(p, f, labeled=lb),
        )


# --- 행 반복 표(row_repeats) -------------------------------------------------
def _make_tag_row(data_tr, tag: str):
    """데이터 행 XML을 복제해 모든 셀을 비우고 첫 셀에 `{%tr ...%}` 태그만 넣는다.

    셀 수·gridSpan을 데이터 행과 동일하게 유지해야 표가 깨지지 않으므로 deepcopy를 쓴다.
    """
    new = copy.deepcopy(data_tr)
    for tc in new.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
    first_p = new.findall(qn("w:tc"))[0].find(qn("w:p"))
    run = first_p.makeelement(qn("w:r"), {})
    wt = first_p.makeelement(qn("w:t"), {})
    wt.text = tag
    run.append(wt)
    first_p.append(run)
    return new


def _set_cell_token(cell, token: str) -> None:
    """표 셀의 첫 문단을 비우고 토큰 하나만 넣는다(데이터 행 컬럼 채움용).

    지우기 전에 그 칸의 글꼴을 읽어 새 런에 그대로 입힌다(표 서식 유지).
    """
    p = cell.paragraphs[0]
    font = read_font(p)
    _clear_runs(p)
    run = p.add_run(token)
    apply_font(run, font)
    mark_font(p, font)


def _apply_row_repeats(doc, mapping) -> None:
    """`row_repeats`: 리스트형 field를 `{%tr%}` 3행 구조로 표에 펼친다(컬럼별 채움).

    데이터 행의 지정 컬럼에 하위필드 토큰을 넣고, 그 앞/뒤에 for·endfor 태그 행을
    삽입한다. 렌더 시 docxtpl가 데이터 행을 항목 수만큼 반복하고 태그 행은 삭제한다.
    """
    default_ti = mapping.get("table", 0)
    for entry in mapping.get("row_repeats", []):
        field = _require(entry, "field", "row_repeats")
        if field not in _ROW_REPEAT:
            raise ValueError(
                f"row_repeats의 field '{field}'는 행 반복 대상이 아닙니다. "
                "허용: " + ", ".join(sorted(_ROW_REPEAT))
            )
        spec = _ROW_REPEAT[field]
        ti = entry.get("table", default_ti)
        table = _at(doc.tables, ti, "표 인덱스", "문서의 표")
        row = _require(entry, "row", "row_repeats")
        data_row = _at(table.rows, row, "행", f"표 {ti}의 행")
        row_cells = data_row.cells
        for subfield, col_idx in _require(entry, "cols", "row_repeats").items():
            if subfield not in spec["col_tokens"]:
                raise ValueError(
                    f"row_repeats field '{field}'에 없는 하위필드 '{subfield}'. "
                    "허용: " + ", ".join(sorted(spec["col_tokens"]))
                )
            cell = _at(row_cells, col_idx, "열", f"행 {row}의 열")
            _check_not_shaded(cell, entry, f"row_repeats[{row},{col_idx}]")
            _set_cell_token(cell, spec["col_tokens"][subfield])
        data_tr = data_row._tr
        for_tag = f"{{%tr for {spec['var']} in {spec['iter']} %}}"
        data_tr.addprevious(_make_tag_row(data_tr, for_tag))
        data_tr.addnext(_make_tag_row(data_tr, "{%tr endfor %}"))


def _collect_drop_rows(doc, mapping) -> list:
    """`drop_rows`가 가리키는 행 요소를 **원본 인덱스 기준으로 먼저** 확보한다.

    양식이 예시 행을 여러 개 깔아 두는 경우(결정사항 4행 등), 데이터보다 남는 행은
    지워야 "내용을 작성하세요."가 남지 않는다. 다만 `row_repeats`가 태그 행을
    끼워 넣어 인덱스가 밀리므로, **바꾸기 전에** 행 요소를 잡아 두고 맨 마지막에
    지운다(인덱스가 아니라 요소 참조라 밀림과 무관).
    """
    doomed = []
    for entry in mapping.get("drop_rows", []):
        ti = entry.get("table", mapping.get("table", 0))
        table = _at(doc.tables, ti, "표 인덱스", "문서의 표")
        for r in _require(entry, "rows", "drop_rows"):
            doomed.append(_at(table.rows, r, "행", f"표 {ti}의 행")._tr)
    return doomed


def apply_mapping(template_path: str, mapping: dict, out_path: str) -> None:
    """mapping대로 양식 복사본에 토큰을 삽입해 out_path에 저장한다.

    `fills`(표 셀)·`paragraphs`(본문 문단)·`row_repeats`(행 반복 표) 중
    있는 것만 적용한다.
    """
    resolved = resolve_input_path(template_path, must_exist=True)
    doc = Document(str(resolved))
    doomed = _collect_drop_rows(doc, mapping)  # 바꾸기 전 인덱스로 확보
    _apply_table_fills(doc, mapping)
    _apply_paragraph_fills(doc, mapping)
    _apply_row_repeats(doc, mapping)
    for tr in doomed:  # 남는 예시 행 제거(맨 마지막)
        tr.getparent().remove(tr)
    doc.save(out_path)


def main() -> None:
    template_path, mapping_path, out_path = sys.argv[1], sys.argv[2], sys.argv[3]
    mapping = json.loads(Path(mapping_path).read_text(encoding="utf-8"))
    apply_mapping(template_path, mapping, out_path)
    print(f"토큰 삽입 완료(토큰화 양식): {out_path}")


if __name__ == "__main__":
    main()
