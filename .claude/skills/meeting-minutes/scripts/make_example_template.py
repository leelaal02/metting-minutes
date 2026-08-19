"""예시 템플릿 생성기: templates/example-template.docx를 코드로 재현한다.

바이너리 docx를 불투명하게 커밋하지 않기 위해, 예시 양식을 이 스크립트로
언제든 재생성한다. 사용자는 산출된 파일을 열어 자기 양식으로 편집하거나
토큰만 복사해 쓴다.
"""
from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "templates" / "example-template.docx"


def build_example() -> Document:
    """10개 항목 토큰을 모두 담은 예시 템플릿 Document를 만든다."""
    doc = Document()
    # 표시자 위치는 자유롭게 편집 가능. 아래는 안내(브레이스/퍼센트 미포함).
    doc.add_paragraph(
        "아래 표시자 위치를 원하는 대로 편집하세요. "
        "실행 항목 표의 행은 항목 수만큼 자동으로 늘어납니다."
    )

    doc.add_heading("{{ title }}", level=0)
    doc.add_paragraph("일시: {{ date }}")
    doc.add_paragraph("참석자: {{ attendees_joined }}")

    doc.add_heading("회의 목적", level=1)
    doc.add_paragraph("{{ purpose }}")

    doc.add_heading("논의 내용", level=1)
    doc.add_paragraph("{%p for d in discussion %}")
    doc.add_paragraph("{{ d.topic }}")
    doc.add_paragraph("{%p for p in d.points %}")
    doc.add_paragraph("- {{ p }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_heading("결정 사항", level=1)
    doc.add_paragraph("{%p for x in decisions %}")
    doc.add_paragraph("- {{ x }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_heading("미결 사항", level=1)
    doc.add_paragraph("{%p for x in open_issues %}")
    doc.add_paragraph("- {{ x }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_heading("실행 항목", level=1)
    # 3행 구조: {%tr for%} 행 / 데이터 행 / {%tr endfor%} 행.
    # docxtpl가 for·endfor 행을 삭제하고 사이 데이터 행을 항목 수만큼 반복한다.
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "할 일", "담당자", "기한"
    table.rows[1].cells[0].text = "{%tr for a in action_items %}"
    data = table.rows[2].cells
    data[0].text = "{{ a.task }}"
    data[1].text = "{{ a.owner }}"
    data[2].text = "{{ a.due }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"

    doc.add_heading("다음 회의", level=1)
    doc.add_paragraph("{{ next_meeting }}")

    doc.add_heading("기타·특이사항", level=1)
    doc.add_paragraph("{%p for n in notes %}")
    doc.add_paragraph("- {{ n }}")
    doc.add_paragraph("{%p endfor %}")
    return doc


def save_example(path: str) -> None:
    build_example().save(path)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    save_example(str(OUT))
    print(f"예시 템플릿 생성 완료: {OUT}")


if __name__ == "__main__":
    main()
