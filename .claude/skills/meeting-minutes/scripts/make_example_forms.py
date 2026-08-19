"""빈 양식(토큰 없는) 예시 생성기: 자동 매핑 경로를 시연·검증하는 에셋을 만든다.

[4] 자동 매핑은 사용자가 준 **토큰이 없는 빈 서식**을 스킬이 알아서 채우는 경로다.
그런데 커밋된 예시가 토큰이 이미 든 example-template.docx뿐이라, 이 경로는 문서로만
존재했다. 여기서 만드는 두 파일이 그 빈틈을 메운다:

- templates/example-blank-table-form.docx    — 표 기반 공식 양식(라벨만, 값칸은 빈칸)
- templates/example-blank-paragraph-form.docx — 문단 기반 공문형 양식(개요·본문 문단)

바이너리 docx를 불투명하게 커밋하지 않도록 언제든 이 스크립트로 재생성한다.
사용자는 산출물을 열어 자기 양식을 만들 때 출발점으로 삼거나, 스킬이 자동 매핑을
어떻게 채우는지 확인하는 데 쓸 수 있다. tests/test_example_forms.py가 이 두 파일에
대해 inspect→apply→render 왕복을 회귀 검증한다.
"""
from pathlib import Path

from docx import Document

TEMPLATES = Path(__file__).resolve().parent.parent / "templates"
TABLE_OUT = TEMPLATES / "example-blank-table-form.docx"
PARA_OUT = TEMPLATES / "example-blank-paragraph-form.docx"


def build_table_form() -> Document:
    """라벨 칸만 채운 빈 표 양식(값 칸은 비워 둔다).

    KISA류 공식 서식을 본떠, 왼쪽에 항목 라벨·오른쪽에 값 칸을 둔 2열 표.
    논의/결정/실행/기타는 라벨 행 아래 넓은 빈칸 행에 담기도록 설계했다.
    """
    doc = Document()
    doc.add_heading("회      의      록", level=0)
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    labels = [
        "제  목",   # row 0 → title
        "일  시",   # row 1 → date
        "참 석 자",  # row 2 → attendees
        "회의 목적",  # row 3 → purpose
        "논의 내용",  # row 4 (아래 빈칸 행에 block)
        "결정 사항",  # row 5
        "미결 사항",  # row 6 → open_issues
        "실행 항목",  # row 7
        "다음 회의",  # row 8 → next_meeting
    ]
    for r, label in enumerate(labels):
        table.rows[r].cells[0].text = label
        # 값 칸(col 1)은 비워 둔다 — 자동 매핑이 여기에 토큰을 넣는다.
    return doc


def build_paragraph_form() -> Document:
    """개요/회의내용 헤딩과 'ㅇ' 항목 문단으로 이뤄진 빈 공문형 양식(표 없음)."""
    doc = Document()
    doc.add_heading("회 의 결 과", level=0)          # 0
    doc.add_heading("1. 개요", level=1)               # 1
    doc.add_paragraph("ㅇ 회의명 :")                   # 2 → title
    doc.add_paragraph("ㅇ 일   시 :")                  # 3 → date
    doc.add_paragraph("ㅇ 참 석 자 :")                 # 4 → attendees
    doc.add_paragraph("ㅇ 목   적 :")                  # 5 → purpose
    doc.add_heading("2. 회의 내용", level=1)           # 6
    doc.add_paragraph("ㅇ")                            # 7 → discussion (block)
    doc.add_heading("3. 결정 사항 및 조치", level=1)   # 8
    doc.add_paragraph("ㅇ")                            # 9 → decisions + action_items (block)
    doc.add_heading("4. 기타", level=1)                # 10
    doc.add_paragraph("ㅇ")                            # 11 → next_meeting + notes (block)
    return doc


def main() -> None:
    TEMPLATES.mkdir(parents=True, exist_ok=True)
    build_table_form().save(str(TABLE_OUT))
    build_paragraph_form().save(str(PARA_OUT))
    print(f"빈 표 양식 생성: {TABLE_OUT}")
    print(f"빈 문단 양식 생성: {PARA_OUT}")


if __name__ == "__main__":
    main()
