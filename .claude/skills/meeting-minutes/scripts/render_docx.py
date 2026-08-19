"""[4] docx 렌더러: minutes.json → .docx (python-docx)."""
import sys

from docx import Document

from validate import load_minutes


def render_docx(data: dict, out_path: str) -> None:
    doc = Document()
    doc.add_heading(data["title"], level=0)
    if data.get("date"):
        doc.add_paragraph(f"일시: {data['date']}")

    doc.add_heading("참석자", level=1)
    if data["attendees"]:
        for a in data["attendees"]:
            doc.add_paragraph(a, style="List Bullet")
    else:
        doc.add_paragraph("(없음)")

    doc.add_heading("회의 목적", level=1)
    doc.add_paragraph(data["purpose"] or "(없음)")

    doc.add_heading("논의 내용", level=1)
    if data["discussion"]:
        for item in data["discussion"]:
            doc.add_heading(item["topic"], level=2)
            for p in item["points"]:
                doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("(없음)")

    doc.add_heading("결정 사항", level=1)
    if data["decisions"]:
        for d in data["decisions"]:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("(없음)")

    doc.add_heading("미결 사항", level=1)
    if data["open_issues"]:
        for o in data["open_issues"]:
            doc.add_paragraph(o, style="List Bullet")
    else:
        doc.add_paragraph("(없음)")

    doc.add_heading("실행 항목 Action Items", level=1)
    if data["action_items"]:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "할 일", "담당자", "기한"
        for a in data["action_items"]:
            row = table.add_row().cells
            row[0].text = a["task"]
            row[1].text = a["owner"] or "-"
            row[2].text = a["due"] or "-"
    else:
        doc.add_paragraph("(없음)")

    doc.add_heading("다음 회의 일정", level=1)
    doc.add_paragraph(data["next_meeting"] or "(미정)")

    doc.add_heading("기타·특이사항", level=1)
    if data["notes"]:
        for n in data["notes"]:
            doc.add_paragraph(n, style="List Bullet")
    else:
        doc.add_paragraph("(없음)")

    # 표 행이 페이지 끝에서 자동으로 나뉘도록 보정한 뒤 저장.
    from docx_postprocess import allow_rows_to_break
    allow_rows_to_break(doc)
    doc.save(out_path)


def main() -> None:
    json_path, out_path = sys.argv[1], sys.argv[2]
    data = load_minutes(json_path)
    render_docx(data, out_path)
    print(f"docx 생성 완료: {out_path}")


if __name__ == "__main__":
    main()
