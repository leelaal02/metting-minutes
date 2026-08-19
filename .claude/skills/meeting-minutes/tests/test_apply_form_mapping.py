"""apply_form_mapping.py 테스트: 토큰 블록 라이브러리 + 삽입 + 엔드투엔드."""
import json
from pathlib import Path

import pytest
from docx import Document

from apply_form_mapping import apply_mapping, block_lines, inline_tokens

SAMPLE = Path(__file__).resolve().parent.parent / "examples" / "sample_minutes.json"


def _sample():
    return json.loads(SAMPLE.read_text(encoding="utf-8"))


def _make_form(path, rows=4, cols=2):
    """빈 표 서식(라벨 칸 포함)을 만든다."""
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    table.cell(1, 0).text = "제 목 :"
    table.cell(2, 0).text = "일 시 :"
    doc.save(str(path))


def _cell_texts(out_path, table=0, row=0, col=0):
    t = Document(str(out_path)).tables[table]
    return [p.text for p in t.rows[row].cells[col].paragraphs]


def _make_shaded_form(path):
    """회색 라벨 행 + 흰 값 행 구조의 양식(누리미디어 형태)."""
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    label = table.cell(0, 0)
    label.text = "회 의 안 건"
    shd = label._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F3F3F3")
    label._tc.get_or_add_tcPr().append(shd)
    doc.save(str(path))


# --- 토큰 블록 라이브러리 단위 테스트 ---------------------------------------
def test_inline_tokens_join():
    # 자동매핑 스칼라는 RichText 슬롯이므로 {{r *_rt }}로 생성한다.
    assert inline_tokens(["title"]) == "{{r title_rt }}"
    assert inline_tokens(["date"]) == "{{r date_rt }}"
    assert inline_tokens(["attendees"]) == "{{r attendees_rt }}"


def test_inline_rejects_list_field():
    with pytest.raises(ValueError, match="목록형"):
        inline_tokens(["discussion"])


def test_block_single_field_no_section_label():
    lines = block_lines(["decisions"])
    assert lines == ["{%p for x in decisions %}", " - {{ x }}", "{%p endfor %}"]
    assert not any(line.startswith("[결정 사항]") for line in lines)


def test_block_single_scalar_is_value_token():
    assert block_lines(["purpose"]) == ["{{r purpose_rt }}"]


def test_block_multi_field_adds_section_labels():
    lines = block_lines(["purpose", "next_meeting"])
    assert lines == ["[회의 목적] {{r purpose_rt }}", "[다음 회의] {{r next_meeting_rt }}"]


def test_block_discussion_uses_numbered_bold_topic():
    lines = block_lines(["discussion"])
    assert lines[0] == "{%p for d in discussion_rt %}"
    assert "{{r d.topic_rt }}" in lines
    assert " - {{ p }}" in lines


def test_block_action_items_owner_due_are_richtext_slots():
    lines = block_lines(["action_items"])
    assert lines[0] == "{%p for a in action_items_rt %}"
    body = " - {{ a.task }} (담당: {{r a.owner_rt }} / 기한: {{r a.due_rt }})"
    assert body in lines


def test_block_multi_list_field_prepends_label_line():
    lines = block_lines(["decisions", "notes"])
    assert lines[0] == "[결정 사항]"
    assert "{%p for x in decisions %}" in lines
    assert "[기타·특이사항]" in lines


def test_unknown_field_raises():
    with pytest.raises(ValueError, match="알 수 없는 field"):
        block_lines(["bogus"])


# --- todo / literal 모드 -----------------------------------------------------
def test_todo_mode_inserts_richtext_todo_token(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(2,0)="일 시 :", 나머지 빈칸
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},   # 빈칸 → todo만
        {"row": 2, "col": 0, "mode": "todo"},   # 라벨칸 → 라벨 뒤 append
    ]}, str(out))
    assert "{{r todo }}" in _cell_texts(out, row=3, col=1)[0]
    label = _cell_texts(out, row=2, col=0)[0]
    assert "일 시 :" in label and "{{r todo }}" in label


def test_todo_mode_needs_no_fields(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    # fields 없이도 동작(예외 없음)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},
    ]}, str(tmp_path / "out.docx"))


def test_todo_mode_end_to_end_red_bold_placeholder(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    cell = Document(str(final)).tables[0].rows[3].cells[1]
    runs = [(r.text, r.bold, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None)
            for p in cell.paragraphs for r in p.runs if r.text]
    assert any(t == "입력필요" and b and c == "FF0000" for t, b, c in runs)


def test_literal_mode_inserts_given_text(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(3,0) 빈칸
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "literal", "text": "서울 본사 3층"},
    ]}, str(out))
    assert "서울 본사 3층" in _cell_texts(out, row=3, col=0)[0]


def test_literal_mode_requires_text_key(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="text"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 3, "col": 0, "mode": "literal"},
        ]}, str(tmp_path / "out.docx"))


# --- row_repeats (행 반복 표) ------------------------------------------------
def _make_action_table_form(path, merge_task=False):
    """헤더 1행 + 빈 데이터행 1행짜리 실행항목 표 양식."""
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    h = t.rows[0].cells
    h[0].text, h[1].text, h[2].text = "할 일", "담당자", "기한"
    if merge_task:
        r = t.rows[1]
        r.cells[0].merge(r.cells[1])  # 담당자 자리까지 병합(gridSpan)
    doc.save(str(path))


def test_row_repeats_builds_tr_structure(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_action_table_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "owner": 1, "due": 2}},
    ]}, str(out))
    t = Document(str(out)).tables[0]
    assert len(t.rows) == 4  # 헤더 + for행 + 데이터행 + endfor행
    cells = [c.text for row in t.rows for c in row.cells]
    assert any("{%tr for a in action_items_rt %}" in x for x in cells)
    assert any("{%tr endfor %}" in x for x in cells)
    assert "{{ a.task }}" in cells
    assert "{{r a.owner_rt }}" in cells
    assert "{{r a.due_rt }}" in cells


def test_row_repeats_end_to_end_columns_and_todo(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_action_table_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "owner": 1, "due": 2}},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    t = Document(str(final)).tables[0]
    assert len(t.rows) == 3  # 헤더 + 항목 2행
    tasks = [t.rows[i].cells[0].text for i in (1, 2)]
    assert "python-docx 렌더러 PoC 작성" in tasks
    assert "샘플 회의 원문 수집" in tasks
    # due=null 항목의 기한 칸 → 빨간 "입력필요"
    for i in (1, 2):
        if t.rows[i].cells[0].text == "샘플 회의 원문 수집":
            due = t.rows[i].cells[2]
            runs = [(r.text, r.bold, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None)
                    for p in due.paragraphs for r in p.runs if r.text]
            assert any(tx == "입력필요" and b and c == "FF0000" for tx, b, c in runs)


def test_row_repeats_preserves_gridspan(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_action_table_form(tpl, merge_task=True)  # task가 c0~c1 병합
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "due": 2}},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    t = Document(str(final)).tables[0]
    assert len(t.rows) == 3  # 헤더 + 2행 (병합 있어도 정상 반복)
    # 각 데이터행에서 병합으로 인해 distinct 셀은 2개(task 병합 + 기한)
    for i in (1, 2):
        distinct = {id(c._tc) for c in t.rows[i].cells}
        assert len(distinct) == 2


def test_row_repeats_unknown_field_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_action_table_form(tpl)
    with pytest.raises(ValueError, match="row_repeats"):
        apply_mapping(str(tpl), {"table": 0, "row_repeats": [
            {"row": 1, "field": "bogus", "cols": {"task": 0}},
        ]}, str(tmp_path / "out.docx"))


# --- 삽입 테스트 -------------------------------------------------------------
def test_inline_appends_after_label(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "inline", "fields": ["title"]},
    ]}, str(out))
    text = _cell_texts(out, row=1, col=0)[0]
    assert "제 목 :" in text            # 라벨 보존
    assert "{{r title_rt }}" in text    # RichText 토큰이 라벨 뒤에


def test_block_single_fills_cell(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    lines = _cell_texts(out, row=3, col=0)
    assert "{%p for x in decisions %}" in lines
    assert " - {{ x }}" in lines


def test_block_multi_keeps_order_and_labels(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["purpose", "next_meeting"]},
    ]}, str(out))
    lines = _cell_texts(out, row=3, col=0)
    assert lines[0] == "[회의 목적] {{r purpose_rt }}"
    assert lines[1] == "[다음 회의] {{r next_meeting_rt }}"


def test_block_preserves_labeled_cell(tmp_path):
    """라벨 칸("제 목 :")에 block을 넣어도 라벨을 지우지 않고 그 뒤에 블록을 붙인다."""
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(1,0) == "제 목 :"
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    lines = _cell_texts(out, row=1, col=0)
    assert lines[0] == "제 목 :"                       # 라벨 보존(첫 문단 유지)
    assert "{%p for x in decisions %}" in lines        # 블록은 라벨 뒤에
    assert lines.index("{%p for x in decisions %}") > 0


def test_shaded_cell_rejected(tmp_path):
    """색칠된 라벨 칸에 값을 넣으려 하면 실패한다 — 조용한 오배치 차단."""
    form = tmp_path / "form.docx"
    _make_shaded_form(form)
    mapping = {"table": 0, "fills": [
        {"row": 0, "col": 0, "mode": "block", "fields": ["purpose"]},
    ]}
    with pytest.raises(ValueError, match="배경색"):
        apply_mapping(str(form), mapping, str(tmp_path / "out.docx"))


def test_shaded_cell_allowed_with_flag(tmp_path):
    """allow_shaded로 명시하면 색칠 칸에도 넣을 수 있다(사용자가 요청한 경우만 쓰는 탈출구)."""
    form = tmp_path / "form.docx"
    _make_shaded_form(form)
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 0, "col": 0, "mode": "inline", "fields": ["title"],
         "allow_shaded": True},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert "{{r title_rt }}" in _cell_texts(out, row=0, col=0)[0]


def test_white_value_cell_still_allowed(tmp_path):
    """음영 없는 값 칸은 그대로 통과한다(가드가 정상 매핑을 막지 않는다)."""
    form = tmp_path / "form.docx"
    _make_shaded_form(form)
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "block", "fields": ["purpose"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    # 라벨 없는 빈 칸이라 제목("[회의 목적]")이 자동으로 앞에 붙는다.
    assert _cell_texts(out, row=1, col=0) == ["[회의 목적]", "{{r purpose_rt }}"]


def test_bullet_only_slot_gets_auto_title(tmp_path):
    """"ㅇ"처럼 글머리만 있는 자리에는 제목이 자동으로 붙는다 — 없으면 정체불명."""
    doc = Document()
    doc.add_paragraph("ㅇ")
    form = tmp_path / "bullet.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["decisions"]}]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).paragraphs[0].text == "ㅇ [결정 사항]"


def test_labeled_slot_gets_no_auto_title(tmp_path):
    """양식에 이미 라벨이 있으면 제목을 덧붙이지 않는다(중복 방지)."""
    doc = Document()
    doc.add_paragraph("결정 사항:")
    form = tmp_path / "labeled.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["decisions"]}]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).paragraphs[0].text == "결정 사항:"


def test_numbered_bullet_slot_is_not_a_label(tmp_path):
    """"1." 같은 번호만 있는 자리도 라벨이 아니다 — 제목이 붙는다."""
    doc = Document()
    doc.add_paragraph("1.")
    form = tmp_path / "numbered_label.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["notes"]}]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).paragraphs[0].text == "1. [기타·특이사항]"


def test_multi_field_block_has_no_duplicate_title(tmp_path):
    """field가 둘 이상이면 block_lines가 이미 라벨을 붙이므로 자동 제목은 생략."""
    doc = Document()
    doc.add_paragraph("ㅇ")
    form = tmp_path / "multi.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [
        {"para": 0, "mode": "block", "fields": ["decisions", "notes"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    result = Document(str(out))
    assert result.paragraphs[0].text == "ㅇ"
    texts = [p.text for p in result.paragraphs]
    assert texts.count("[결정 사항]") == 1


def test_item_marker_avoids_template_marker(tmp_path):
    """양식이 "-"를 쓰면 하위 항목 기호를 "·"로 바꾼다 — 같은 계층 기호 중복 방지."""
    doc = Document()
    doc.add_paragraph("-")
    form = tmp_path / "dash.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["decisions"]}]}
    apply_mapping(str(form), mapping, str(out))
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any(t.strip().startswith("· {{") for t in texts)
    assert not any(t.strip().startswith("- {{") for t in texts)


def test_topic_marker_switches_when_template_numbers(tmp_path):
    """양식이 번호를 쓰고 그 번호가 남는 자리면 주제는 넘버링 대신 기호를 쓴다."""
    doc = Document()
    doc.add_paragraph("1. 회의 내용")  # 라벨 + 번호가 그대로 남는 자리
    form = tmp_path / "num.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["discussion"]}]}
    apply_mapping(str(form), mapping, str(out))
    texts = [p.text.strip() for p in Document(str(out)).paragraphs]
    assert "□ {{r d.topic_plain_rt }}" in texts   # 기호 + 번호 없는 주제
    assert "{{r d.topic_rt }}" not in texts       # 넘버링 토큰은 안 씀


def test_open_issues_block_is_supported(tmp_path):
    """양식의 "미결 사항" 칸에 open_issues를 넣을 수 있다(todo로 비우지 않는다)."""
    doc = Document()
    doc.add_paragraph("미결 사항:")
    form = tmp_path / "open.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["open_issues"]}]}
    apply_mapping(str(form), mapping, str(out))
    texts = [p.text.strip() for p in Document(str(out)).paragraphs]
    assert "{%p for x in open_issues %}" in texts
    assert "{{r todo }}" not in " ".join(texts)  # 데이터 자리를 "입력필요"로 덮지 않음


def test_attendee_count_gap_filled(tmp_path):
    """양식이 비워 둔 "총  명"에 참석인원 수 토큰이 끼워진다 — 명단은 뒤에 이어 붙는다.

    워드가 한 문장을 여러 런으로 쪼개 둔 실제 양식(`총`/`  `/`명`)을 그대로 재현한다
    — 런 하나만 보는 구현은 이 빈칸을 놓친다.
    """
    doc = Document()
    p = doc.add_paragraph()
    for chunk in ["ㅇ", " (", "참석인원", ") ", "총", "  ", "명", " ", "참석"]:
        p.add_run(chunk)
    form = tmp_path / "count.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "inline", "fields": ["attendees"]}]}
    apply_mapping(str(form), mapping, str(out))
    text = Document(str(out)).paragraphs[0].text
    assert "총 {{ attendee_count }}명 참석" in text
    assert text.rstrip().endswith("{{r attendees_rt }}")


def test_attendee_count_gap_untouched_without_gap(tmp_path):
    """빈칸이 없는 자리는 건드리지 않는다 — "총명"·"참석자:"에 수를 밀어 넣지 않는다."""
    doc = Document()
    doc.add_paragraph("참 석 자 :")
    form = tmp_path / "nogap.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "inline", "fields": ["attendees"]}]}
    apply_mapping(str(form), mapping, str(out))
    text = Document(str(out)).paragraphs[0].text
    assert "attendee_count" not in text
    assert text == "참 석 자 : {{r attendees_rt }}"


def test_topic_keeps_numbering_when_template_number_disappears(tmp_path):
    """양식 자동번호가 렌더 시 사라지는 자리면 주제 넘버링("1. 주제")을 그대로 쓴다.

    빈 자동번호 문단에는 `{%p for … %}` 태그가 들어가고 그 문단은 렌더에서 삭제된다
    — 번호가 남지 않으므로 겹칠 일이 없는데 기호(`□`)로 내려가면 넘버링을 헛되게
    버린다. `auto_label: false`로 제목을 끈 자리에서 실제로 이 손해가 났다.
    """
    from docx.oxml.ns import qn

    doc = Document()
    target = doc.add_paragraph("")  # 글자 없는 자동번호 문단 = 값 자리
    pPr = target._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn("w:numPr"), {})
    numId = pPr.makeelement(qn("w:numId"), {})
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    form = tmp_path / "autonum.docx"
    doc.save(str(form))

    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [
        {"para": 0, "mode": "block", "fields": ["discussion"], "auto_label": False},
    ]}
    apply_mapping(str(form), mapping, str(out))
    texts = [p.text.strip() for p in Document(str(out)).paragraphs]
    assert "{{r d.topic_rt }}" in texts             # 넘버링 토큰을 씀
    assert "□ {{r d.topic_plain_rt }}" not in texts  # 기호로 내려가지 않음


def test_marker_ladder_falls_back_when_exhausted(tmp_path):
    """양식이 사다리를 다 쓰면 마지막 기호로 떨어진다(들여쓰기로 계층 유지)."""
    from apply_form_mapping import _ITEM_LADDER, _choose_markers

    assert _choose_markers(set(_ITEM_LADDER) | {"num"})["item"] == _ITEM_LADDER[-1]


def test_nested_slot_indents_one_more_step(tmp_path):
    """양식이 이미 한 계층("ㅇ")을 쓰면 우리 내용은 그 아래로 한 단 더 들어간다."""
    from docx.shared import Cm

    plain = Document()
    plain.add_paragraph("")
    plain_form = tmp_path / "plain.docx"
    plain.save(str(plain_form))
    bulleted = Document()
    bulleted.add_paragraph("ㅇ")
    bullet_form = tmp_path / "bullet.docx"
    bulleted.save(str(bullet_form))

    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["decisions"]}]}
    indents = []
    for src, name in ((plain_form, "a.docx"), (bullet_form, "b.docx")):
        out = tmp_path / name
        apply_mapping(str(src), mapping, str(out))
        item = [p for p in Document(str(out)).paragraphs
                if p.text.strip().startswith(("-", "·"))][0]
        indents.append(item.paragraph_format.left_indent)
    assert abs((indents[1] - indents[0]) - Cm(0.4)) < Cm(0.01)


def test_block_lines_get_hanging_indent(tmp_path):
    """하위 항목("- …")은 들여쓰기 + 내어쓰기 — 줄바꿈돼도 글머리에 정렬된다."""
    from docx.shared import Cm

    form = tmp_path / "form.docx"
    _make_form(form)
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    cell = Document(str(out)).tables[0].rows[3].cells[0]
    item = [p for p in cell.paragraphs if p.text.strip().startswith("-")][0]
    # docx는 들여쓰기를 twip으로 저장해 EMU 왕복 시 오차가 생긴다 → 근사 비교.
    assert abs(item.paragraph_format.left_indent - Cm(1.1)) < Cm(0.01)
    assert abs(item.paragraph_format.first_line_indent - Cm(-0.35)) < Cm(0.01)


def test_section_label_is_bold_with_wide_space(tmp_path):
    """섹션 제목은 굵게 + 넓은 위 여백 — 하위 계층과 구분된다."""
    from docx.shared import Cm, Pt

    form = tmp_path / "form.docx"
    _make_form(form)
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions", "notes"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    cell = Document(str(out)).tables[0].rows[3].cells[0]
    label = [p for p in cell.paragraphs if p.text.strip().startswith("[")][0]
    assert label.paragraph_format.space_before == Pt(10)
    assert label.paragraph_format.left_indent == Cm(0)
    assert all(r.bold for r in label.runs)


def test_three_levels_have_distinct_indent(tmp_path):
    """섹션 제목 / 소제목 / 하위 항목이 서로 다른 들여쓰기를 갖는다.

    계층마다 기호(`[…]` / `1.` / `-`)가 다르고 들여쓰기도 벌어져야 읽힌다.
    """
    form = tmp_path / "form.docx"
    _make_form(form)
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["discussion", "notes"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    cell = Document(str(out)).tables[0].rows[3].cells[0]
    by_mark = {}
    for p in cell.paragraphs:
        text = p.text.strip()
        if text.startswith("["):
            by_mark.setdefault("section", p.paragraph_format.left_indent)
        elif text.startswith("{{r d.topic_rt"):
            by_mark.setdefault("topic", p.paragraph_format.left_indent)
        elif text.startswith("-"):
            by_mark.setdefault("item", p.paragraph_format.left_indent)
    assert by_mark["section"] < by_mark["topic"] < by_mark["item"]


def test_inserted_paragraph_inherits_format_without_numbering(tmp_path):
    """삽입 문단은 원본 서식을 물려받되 번호(numPr)는 물려받지 않는다.

    numPr까지 복사하면 목록 각 줄에 "1. 2. 3."이 덧붙는다.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    target = doc.add_paragraph("")
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = target._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn("w:numPr"), {})
    numId = pPr.makeelement(qn("w:numId"), {})
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    form = tmp_path / "para.docx"
    doc.save(str(form))

    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [{"para": 0, "mode": "block", "fields": ["decisions"]}]}
    apply_mapping(str(form), mapping, str(out))

    result = Document(str(out))
    items = [p for p in result.paragraphs if p.text.strip().startswith("-")]
    assert items, "블록이 삽입되지 않았다"
    for p in items:
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER  # 서식은 상속
        assert p._p.find(qn("w:pPr")).find(qn("w:numPr")) is None  # 번호는 제외


def test_placeholder_text_is_replaced(tmp_path):
    """양식의 예시 문구("내용을 작성하세요.")는 지우고 값을 넣는다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "특이사항"
    cell = table.cell(0, 1)
    cell.text = "내용을 작성하세요."
    cell.add_paragraph("본문 글자크기는 8.5pt를 권장합니다.")
    form = tmp_path / "ph.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 0, "col": 1, "mode": "block", "fields": ["notes"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    text = Document(str(out)).tables[0].cell(0, 1).text
    assert "작성하세요" not in text and "권장합니다" not in text
    assert "{%p for n in notes %}" in text


def test_filler_placeholder_is_replaced(tmp_path):
    """"2025년 00월 00일"·"OO팀 OOO" 같은 빈칸 표기도 예시로 본다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "회의일시"
    table.cell(0, 1).text = "2025년 00월 00일"
    form = tmp_path / "filler.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 0, "col": 1, "mode": "inline", "fields": ["date"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).tables[0].cell(0, 1).text.strip() == "{{r date_rt }}"


def test_real_label_is_not_treated_as_placeholder(tmp_path):
    """평범한 라벨("회의일시")은 예시가 아니므로 보존된다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "회의일시"
    form = tmp_path / "label.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [
        {"row": 0, "col": 0, "mode": "inline", "fields": ["date"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).tables[0].cell(0, 0).text.startswith("회의일시")


def test_duplicate_label_hint_is_cleared(tmp_path):
    """값 칸이 라벨과 같은 글자를 담고 있으면("부서 | 부서") 힌트로 보고 지운다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "부서"
    table.cell(0, 1).text = "부서"
    form = tmp_path / "hint.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"table": 0, "fills": [{"row": 0, "col": 1, "mode": "todo"}]}
    apply_mapping(str(form), mapping, str(out))
    assert Document(str(out)).tables[0].cell(0, 1).text.strip() == "{{r todo }}"


def test_fills_can_target_different_tables(tmp_path):
    """표가 여러 개인 양식도 항목별 table 키로 한 번에 채운다."""
    doc = Document()
    doc.add_table(rows=1, cols=1)
    doc.add_table(rows=1, cols=1)
    form = tmp_path / "two.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"fills": [
        {"table": 0, "row": 0, "col": 0, "mode": "inline", "fields": ["title"]},
        {"table": 1, "row": 0, "col": 0, "mode": "inline", "fields": ["date"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    result = Document(str(out))
    assert "{{r title_rt }}" in result.tables[0].cell(0, 0).text
    assert "{{r date_rt }}" in result.tables[1].cell(0, 0).text


def test_drop_rows_removes_leftover_example_rows(tmp_path):
    """양식이 깔아 둔 예시 행은 drop_rows로 지운다(데이터보다 남는 행)."""
    doc = Document()
    table = doc.add_table(rows=4, cols=1)
    for r in range(1, 4):
        table.cell(r, 0).text = "내용을 작성하세요."
    form = tmp_path / "rows.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {
        "table": 0,
        "fills": [{"row": 1, "col": 0, "mode": "block", "fields": ["decisions"]}],
        "drop_rows": [{"table": 0, "rows": [2, 3]}],
    }
    apply_mapping(str(form), mapping, str(out))
    result = Document(str(out)).tables[0]
    assert len(result.rows) == 2
    assert "작성하세요" not in result.cell(1, 0).text


def test_out_of_range_row_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl, rows=4)
    with pytest.raises(IndexError, match="행"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 99, "col": 0, "mode": "inline", "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_out_of_range_table_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(IndexError, match="표 인덱스"):
        apply_mapping(str(tpl), {"table": 5, "fills": [
            {"row": 0, "col": 0, "mode": "inline", "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_missing_fill_key_raises_friendly_error(tmp_path):
    # mode 키가 빠진 fills 항목 → 원시 KeyError 대신 안내 메시지(ValueError)
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="필수 키 'mode'"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 1, "col": 0, "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_missing_paragraph_key_raises_friendly_error(tmp_path):
    tpl = tmp_path / "pform.docx"
    _make_para_form(tpl)
    with pytest.raises(ValueError, match="필수 키 'fields'"):
        apply_mapping(str(tpl), {"paragraphs": [
            {"para": 1, "mode": "inline"},
        ]}, str(tmp_path / "out.docx"))


def test_unknown_field_in_apply_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="알 수 없는 field"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 3, "col": 0, "mode": "block", "fields": ["nope"]},
        ]}, str(tmp_path / "out.docx"))


# --- 문단 기반 양식 삽입 -----------------------------------------------------
def _make_para_form(path):
    """개요/참석자/회의내용 문단 양식(표 없음)을 만든다."""
    doc = Document()
    doc.add_paragraph("개요")          # 0
    doc.add_paragraph("ㅇ (목적)")     # 1
    doc.add_paragraph("ㅇ (일시)")     # 2
    doc.add_paragraph("회의내용")      # 3
    doc.add_paragraph("ㅇ")            # 4
    doc.add_paragraph("ㅇ")            # 5
    doc.save(str(path))


def _para_texts(out_path):
    return [p.text for p in Document(str(out_path)).paragraphs]


def test_paragraph_inline_appends_after_label(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 1, "mode": "inline", "fields": ["purpose"]},
        {"para": 2, "mode": "inline", "fields": ["date"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert "ㅇ (목적) {{r purpose_rt }}" in texts
    assert "ㅇ (일시) {{r date_rt }}" in texts


def test_paragraph_block_inserts_following_paragraphs(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 4, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    texts = _para_texts(out)
    # block 첫 라인은 대상 문단에, 나머지는 바로 뒤 문단으로
    assert "{%p for x in decisions %}" in texts
    assert " - {{ x }}" in texts
    assert "{%p endfor %}" in texts
    # 회의내용 헤딩은 보존되고, 뒤의 다른 ㅇ 문단도 남아 있다
    assert "회의내용" in texts


def test_paragraph_block_multi_field_labels(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 4, "mode": "block", "fields": ["decisions", "notes"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert "[결정 사항]" in texts
    assert "[기타·특이사항]" in texts


def test_block_preserves_labeled_paragraph(tmp_path):
    """라벨 문단에 block을 넣어도 라벨을 지우지 않고 바로 뒤에 블록을 삽입한다."""
    doc = Document()
    doc.add_paragraph("ㅇ 결정사항")   # 0: 라벨 문단
    doc.add_paragraph("뒤 문단")        # 1: 이후 문단 보존 확인용
    tpl = tmp_path / "pform.docx"
    doc.save(str(tpl))
    out = tmp_path / "out.docx"
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert texts[0] == "ㅇ 결정사항"                    # 라벨 보존(그 자리 유지)
    assert texts[1] == "{%p for x in decisions %}"      # 블록이 라벨 바로 뒤
    assert "뒤 문단" in texts                           # 이후 문단도 보존


def test_paragraph_out_of_range_raises(tmp_path):
    tpl = tmp_path / "pform.docx"
    _make_para_form(tpl)
    with pytest.raises(IndexError, match="문단 인덱스"):
        apply_mapping(str(tpl), {"paragraphs": [
            {"para": 999, "mode": "inline", "fields": ["purpose"]},
        ]}, str(tmp_path / "out.docx"))


def test_mixed_table_and_paragraph_fills(tmp_path):
    """표 셀(title) + 문단(purpose)을 한 번에 채운다."""
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "회의록"
    doc.add_paragraph("ㅇ (목적)")  # 문단 index 0 (표 안 문단은 제외되므로)
    tpl = tmp_path / "mixed.docx"
    doc.save(str(tpl))
    out = tmp_path / "out.docx"
    apply_mapping(str(tpl), {
        "table": 0,
        "fills": [{"row": 0, "col": 0, "mode": "inline", "fields": ["title"]}],
        "paragraphs": [{"para": 0, "mode": "inline", "fields": ["purpose"]}],
    }, str(out))
    doc_out = Document(str(out))
    assert "회의록 {{r title_rt }}" in doc_out.tables[0].rows[0].cells[0].text
    assert "ㅇ (목적) {{r purpose_rt }}" in [p.text for p in doc_out.paragraphs]


def test_paragraph_block_end_to_end(tmp_path):
    """문단 block 삽입 → render_docx_template로 실제 값이 채워지는지."""
    from render_docx_template import render_template

    tpl = tmp_path / "pform.docx"
    tokenized = tmp_path / "pform_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 1, "mode": "inline", "fields": ["purpose"]},
        {"para": 4, "mode": "block", "fields": ["decisions"]},
        {"para": 5, "mode": "block", "fields": ["action_items"]},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    all_text = "\n".join(_para_texts(final))
    assert "docx 변환은 python-docx로 진행" in all_text     # decisions
    assert "python-docx 렌더러 PoC 작성" in all_text         # action_items


# --- 엔드투엔드: apply → render_docx_template ---------------------------------
def test_end_to_end_fills_values(tmp_path):
    from render_docx_template import render_template

    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl, rows=5)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "inline", "fields": ["title"]},
        {"row": 2, "col": 0, "mode": "inline", "fields": ["date"]},
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions"]},
        {"row": 4, "col": 0, "mode": "block", "fields": ["action_items"]},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))

    table = Document(str(final)).tables[0]
    all_text = "\n".join(
        p.text for row in table.rows for cell in row.cells for p in cell.paragraphs
    )
    assert "2026 3분기 제품 로드맵 회의" in all_text     # title inline
    assert "2026-07-16 14:00" in all_text               # date inline
    assert "docx 변환은 python-docx로 진행" in all_text  # decisions 반복
    assert "python-docx 렌더러 PoC 작성" in all_text     # action_items 반복


def test_end_to_end_empty_list_leaves_blank(tmp_path):
    from render_docx_template import render_template

    data = _sample()
    data["notes"] = []
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl, rows=4)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["notes"]},
    ]}, str(tokenized))
    render_template(str(tokenized), data, str(final))

    lines = _cell_texts(final, row=3, col=0)
    # notes가 비면 반복 문단이 생성되지 않아 " - ..." 라인이 없다
    assert not any(line.strip().startswith("-") for line in lines)


# --- 자동 섹션 제목: 옆 칸 라벨 배치 ------------------------------------------
def _make_label_value_form(path):
    """`라벨 칸 | 값 칸` 배치의 표 양식 — 한국 회의록 양식의 표준 구조."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "논의 내용"
    table.cell(1, 0).text = "결정 사항"
    doc.save(str(path))


def test_adjacent_label_cell_suppresses_auto_title(tmp_path):
    """옆 칸이 라벨이면 제목을 붙이지 않는다 — "논의 내용 | [논의 내용]" 중복 방지."""
    form = tmp_path / "labeled_row.docx"
    _make_label_value_form(form)
    out = tmp_path / "out.docx"
    mapping = {"fills": [
        {"row": 0, "col": 1, "mode": "block", "fields": ["discussion"]},
        {"row": 1, "col": 1, "mode": "block", "fields": ["decisions"]},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert not any("[논의 내용]" in t for t in _cell_texts(out, row=0, col=1))
    assert not any("[결정 사항]" in t for t in _cell_texts(out, row=1, col=1))
    # 라벨 칸은 그대로 보존되고 내용 블록은 값 칸에 들어간다.
    assert _cell_texts(out, row=0, col=0) == ["논의 내용"]
    assert any("{%p for d in discussion_rt %}" in t for t in _cell_texts(out, row=0, col=1))


def test_unlabeled_row_still_gets_auto_title(tmp_path):
    """행 전체에 라벨이 없으면 제목은 여전히 자동으로 붙는다(정체불명 방지)."""
    doc = Document()
    doc.add_table(rows=1, cols=2)  # 라벨 없는 빈 표
    form = tmp_path / "bare_row.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"fills": [{"row": 0, "col": 1, "mode": "block", "fields": ["decisions"]}]}
    apply_mapping(str(form), mapping, str(out))
    assert _cell_texts(out, row=0, col=1)[0] == "[결정 사항]"


def test_auto_label_false_turns_title_off(tmp_path):
    """라벨이 위쪽 행에 있는 양식은 `auto_label: false`로 제목을 끈다."""
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "회 의 내 용"  # 라벨 행
    form = tmp_path / "label_above.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"fills": [
        {"row": 1, "col": 0, "mode": "block", "fields": ["discussion"],
         "auto_label": False},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert not any("[논의 내용]" in t for t in _cell_texts(out, row=1, col=0))


def test_auto_label_true_forces_title(tmp_path):
    """옆 칸 라벨이 이 항목과 무관하면 `auto_label: true`로 제목을 강제한다."""
    form = tmp_path / "forced.docx"
    _make_label_value_form(form)
    out = tmp_path / "out.docx"
    mapping = {"fills": [
        {"row": 0, "col": 1, "mode": "block", "fields": ["notes"], "auto_label": True},
    ]}
    apply_mapping(str(form), mapping, str(out))
    assert _cell_texts(out, row=0, col=1)[0] == "[기타·특이사항]"


def test_auto_label_false_works_on_paragraph_slot(tmp_path):
    """본문 문단도 `auto_label: false`로 제목을 끌 수 있다(앞 문단이 라벨인 양식)."""
    doc = Document()
    doc.add_paragraph("회의 내용")  # 라벨 문단
    doc.add_paragraph("")           # 값 자리
    form = tmp_path / "para_label_above.docx"
    doc.save(str(form))
    out = tmp_path / "out.docx"
    mapping = {"paragraphs": [
        {"para": 1, "mode": "block", "fields": ["decisions"], "auto_label": False},
    ]}
    apply_mapping(str(form), mapping, str(out))
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert "[결정 사항]" not in texts
    assert texts[0] == "회의 내용"
