import json
from pathlib import Path

import pytest
from docx import Document

from render_docx_template import build_context

SAMPLE = Path(__file__).resolve().parent.parent / "examples" / "sample_minutes.json"


def _sample():
    return json.loads(SAMPLE.read_text(encoding="utf-8"))


def test_build_context_passes_scalars_and_joins_attendees():
    ctx = build_context(_sample())
    assert ctx["title"] == "2026 3분기 제품 로드맵 회의"
    assert ctx["date"] == "2026-07-16 14:00"
    assert ctx["attendees_joined"] == "김수민, 이정우, 박서연"
    assert ctx["attendees"] == ["김수민", "이정우", "박서연"]


def test_build_context_null_scalars_become_empty_string():
    data = _sample()
    data["date"] = None
    data["purpose"] = None
    data["next_meeting"] = None
    ctx = build_context(data)
    assert ctx["date"] == ""
    assert ctx["purpose"] == ""
    assert ctx["next_meeting"] == ""


def test_build_context_null_owner_due_become_dash():
    ctx = build_context(_sample())
    # 샘플 2번째 action_item은 due=null
    tasks = {a["task"]: a for a in ctx["action_items"]}
    assert tasks["샘플 회의 원문 수집"]["due"] == "-"
    assert tasks["샘플 회의 원문 수집"]["owner"] == "박서연"


def test_build_context_empty_attendees_joined_is_empty():
    data = _sample()
    data["attendees"] = []
    ctx = build_context(data)
    assert ctx["attendees_joined"] == ""
    assert ctx["attendees"] == []


def test_build_context_todo_is_red_bold_richtext():
    from docxtpl import RichText
    ctx = build_context(_sample())
    assert isinstance(ctx["todo"], RichText)
    xml = ctx["todo"].xml
    assert "입력필요" in xml
    assert "FF0000" in xml
    assert "<w:b/>" in xml


def test_build_context_scalar_rt_present_is_plain_value():
    from docxtpl import RichText
    ctx = build_context(_sample())
    assert isinstance(ctx["date_rt"], RichText)
    assert "2026-07-16 14:00" in ctx["date_rt"].xml
    assert "FF0000" not in ctx["date_rt"].xml   # 값 있으면 빨강 아님


def test_build_context_scalar_rt_empty_is_todo():
    data = _sample()
    data["date"] = None
    data["purpose"] = None
    data["next_meeting"] = None
    ctx = build_context(data)
    for key in ("date_rt", "purpose_rt", "next_meeting_rt"):
        assert "입력필요" in ctx[key].xml
        assert "FF0000" in ctx[key].xml


def test_build_context_attendees_rt_empty_is_todo():
    data = _sample()
    data["attendees"] = []
    ctx = build_context(data)
    assert "입력필요" in ctx["attendees_rt"].xml
    assert "FF0000" in ctx["attendees_rt"].xml


def test_build_context_discussion_rt_numbered_bold_and_original_unchanged():
    data = _sample()
    ctx = build_context(data)
    d0 = ctx["discussion_rt"][0]
    assert "1. STT 연동 우선순위" in d0["topic_rt"].xml
    assert "<w:b/>" in d0["topic_rt"].xml
    assert ctx["discussion_rt"][1]["topic_rt"].xml.count("2. 출력 포맷") == 1
    assert d0["points"] == ["실시간 STT는 3분기 범위에서 제외", "배치 STT부터 도입하기로 의견 수렴"]
    # 원본 data는 변형되지 않는다(topic_rt 주입 금지)
    assert "topic_rt" not in data["discussion"][0]


def test_build_context_action_items_rt_owner_due():
    ctx = build_context(_sample())
    items = {a_rt["task"]: a_rt for a_rt in ctx["action_items_rt"]}
    # due=null 인 항목 → due_rt는 todo(빨강), owner는 값
    it = items["샘플 회의 원문 수집"]
    assert "박서연" in it["owner_rt"].xml
    assert "FF0000" not in it["owner_rt"].xml
    assert "입력필요" in it["due_rt"].xml
    assert "FF0000" in it["due_rt"].xml


def test_build_context_keeps_plain_keys_backcompat():
    """평문 키(사용자 {{ }} 경로)는 그대로 유지된다."""
    ctx = build_context(_sample())
    assert ctx["title"] == "2026 3분기 제품 로드맵 회의"
    assert ctx["date"] == "2026-07-16 14:00"
    assert ctx["attendees_joined"] == "김수민, 이정우, 박서연"
    assert {a["task"] for a in ctx["action_items"]} == {"python-docx 렌더러 PoC 작성", "샘플 회의 원문 수집"}


def _make_template(path):
    """토큰 + {%tr%} 표(3행 구조)를 담은 최소 템플릿 docx를 만든다."""
    doc = Document()
    doc.add_paragraph("{{ title }}")
    doc.add_paragraph("일시: {{ date }}")
    doc.add_paragraph("목적: {{ purpose }}")
    doc.add_paragraph("참석자: {{ attendees_joined }}")
    doc.add_paragraph("다음 회의: {{ next_meeting }}")
    # 실행 항목 표: {%tr for%} 행 / 데이터 행 / {%tr endfor%} 행 (3행 구조).
    # docxtpl가 {%tr ...%} 든 행 전체를 삭제하고 for·endfor 사이 행을 반복한다.
    table = doc.add_table(rows=4, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "할 일", "담당자", "기한"
    table.rows[1].cells[0].text = "{%tr for a in action_items %}"
    data = table.rows[2].cells
    data[0].text = "{{ a.task }}"
    data[1].text = "{{ a.owner }}"
    data[2].text = "{{ a.due }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"
    doc.save(str(path))


def _rendered_paragraphs(out_path):
    return [p.text for p in Document(str(out_path)).paragraphs]


def test_simple_tokens_substituted(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), _sample(), str(out))
    texts = _rendered_paragraphs(out)
    assert "2026 3분기 제품 로드맵 회의" in texts
    assert "일시: 2026-07-16 14:00" in texts
    assert "참석자: 김수민, 이정우, 박서연" in texts


def test_action_items_table_rows_repeat(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), _sample(), str(out))
    table = Document(str(out)).tables[0]
    # 헤더 1행 + action_items 2행
    assert len(table.rows) == 3
    body_tasks = [table.rows[i].cells[0].text for i in (1, 2)]
    assert "python-docx 렌더러 PoC 작성" in body_tasks
    assert "샘플 회의 원문 수집" in body_tasks


def test_null_scalar_renders_blank(tmp_path):
    from render_docx_template import render_template
    data = _sample()
    data["next_meeting"] = None
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), data, str(out))
    assert "다음 회의: " in _rendered_paragraphs(out)


def test_null_due_renders_dash(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), _sample(), str(out))
    table = Document(str(out)).tables[0]
    due_cells = [table.rows[i].cells[2].text for i in (1, 2)]
    assert "-" in due_cells


def test_empty_action_items_leaves_header_only(tmp_path):
    from render_docx_template import render_template
    data = _sample()
    data["action_items"] = []
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), data, str(out))
    table = Document(str(out)).tables[0]
    assert len(table.rows) == 1  # 헤더만 남음


def test_non_docx_template_raises(tmp_path):
    from render_docx_template import render_template
    bad = tmp_path / "form.txt"
    bad.write_text("not a docx", encoding="utf-8")
    with pytest.raises(ValueError, match="docx"):
        render_template(str(bad), _sample(), str(tmp_path / "out.docx"))


def test_missing_template_raises(tmp_path):
    from render_docx_template import render_template
    missing = tmp_path / "nope.docx"  # 존재하지 않는 절대경로
    with pytest.raises(FileNotFoundError):
        render_template(str(missing), _sample(), str(tmp_path / "out.docx"))


def test_malformed_token_raises_helpful_error(tmp_path):
    from render_docx_template import render_template
    doc = Document()
    # endfor 없는 for 블록 → Jinja 문법 오류
    doc.add_paragraph("{% for x in decisions %}")
    doc.add_paragraph("- {{ x }}")
    tpl = tmp_path / "bad.docx"
    doc.save(str(tpl))
    with pytest.raises(ValueError, match="문법"):
        render_template(str(tpl), _sample(), str(tmp_path / "out.docx"))


def test_xml_special_chars_in_values_preserved(tmp_path):
    """값에 든 &, <, > 가 렌더 후에도 보존된다(autoescape). 회귀 방지."""
    from render_docx_template import render_template
    data = _sample()
    data["title"] = "R&D <검토> 회의 & 계획"
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    _make_template(tpl)
    render_template(str(tpl), data, str(out))
    assert "R&D <검토> 회의 & 계획" in _rendered_paragraphs(out)


def test_committed_example_template_renders(tmp_path):
    from render_docx_template import render_template
    committed = Path(__file__).resolve().parent.parent / "templates" / "example-template.docx"
    out = tmp_path / "out.docx"
    render_template(str(committed), _sample(), str(out))
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("2026 3분기 제품 로드맵 회의" in t for t in texts)


def test_example_template_renders_with_sample(tmp_path):
    from make_example_template import build_example
    from render_docx_template import render_template
    tpl = tmp_path / "example.docx"
    build_example().save(str(tpl))
    out = tmp_path / "out.docx"
    render_template(str(tpl), _sample(), str(out))
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    # 제목 치환 확인
    assert any("2026 3분기 제품 로드맵 회의" in t for t in texts)
    # 결정 사항 목록 반복 확인
    assert any("docx 변환은 python-docx로 진행" in t for t in texts)
    # 실행 항목 표에 데이터 행 존재
    table = doc.tables[0]
    assert len(table.rows) >= 3  # 헤더 + 2행 이상
