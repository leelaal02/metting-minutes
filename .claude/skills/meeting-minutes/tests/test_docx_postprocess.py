"""docx_postprocess.py 테스트: 표 행 나눔 금지 해제 + 양식 글꼴 물려주기."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docx_postprocess import allow_rows_to_break, inherit_mark_fonts


def _mark_font(paragraph, name="바탕", size=14):
    """문단 부호(w:pPr/w:rPr)에 글꼴 표식을 남긴다 — apply 단계가 하는 일."""
    pPr = paragraph._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    pPr.insert_element_before(rPr, "w:sectPr", "w:pPrChange")
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), name)
    rPr.sz_val = Pt(size)


def test_mark_font_applied_to_bare_run():
    """글꼴이 없는 런(렌더로 새로 생긴 RichText)에 양식 글꼴이 입혀진다."""
    doc = Document()
    p = doc.add_paragraph("주제")
    _mark_font(p)
    assert inherit_mark_fonts(doc) == 1
    rPr = p.runs[0]._r.find(qn("w:rPr"))
    assert rPr.find(qn("w:rFonts")).get(qn("w:eastAsia")) == "바탕"
    assert p.runs[0].font.size == Pt(14)


def test_mark_font_does_not_override_explicit_run_font():
    """런에 이미 지정된 글꼴이 우선 — 표식이 덮어쓰지 않는다."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("주제")
    run.font.size = Pt(20)
    _mark_font(p, size=14)
    inherit_mark_fonts(doc)
    assert run.font.size == Pt(20)


def test_paragraph_without_mark_is_untouched():
    """표식이 없는 문단(양식 원래 내용)은 손대지 않는다."""
    doc = Document()
    p = doc.add_paragraph("양식 원문")
    assert inherit_mark_fonts(doc) == 0
    assert p.runs[0]._r.find(qn("w:rPr")) is None


def _set_cant_split(row):
    """행에 나눔 금지(w:cantSplit)를 건다(양식 템플릿이 해 두는 설정 재현)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _has_cant_split(row) -> bool:
    trPr = row._tr.find(qn("w:trPr"))
    return trPr is not None and trPr.find(qn("w:cantSplit")) is not None


def test_removes_cant_split_from_all_rows(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    for row in table.rows:
        _set_cant_split(row)
    assert all(_has_cant_split(r) for r in table.rows)

    removed = allow_rows_to_break(doc)

    assert removed == 3
    assert not any(_has_cant_split(r) for r in table.rows)


def test_rows_without_cant_split_untouched(tmp_path):
    # 나눔 금지가 없던 행은 그대로(제거할 것 없음 → 0 반환)
    doc = Document()
    doc.add_table(rows=2, cols=2)
    assert allow_rows_to_break(doc) == 0


def test_nested_table_rows_also_fixed(tmp_path):
    # 셀 안 중첩표의 행도 재귀로 처리된다
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    _set_cant_split(outer.rows[0])
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    _set_cant_split(nested.rows[0])

    removed = allow_rows_to_break(doc)

    assert removed == 2
    assert not _has_cant_split(outer.rows[0])
    assert not _has_cant_split(nested.rows[0])


def test_round_trip_persists(tmp_path):
    # 저장 후 다시 열어도 cantSplit가 없어야 한다
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _set_cant_split(table.rows[0])
    allow_rows_to_break(doc)
    p = tmp_path / "out.docx"
    doc.save(str(p))
    reopened = Document(str(p))
    assert not _has_cant_split(reopened.tables[0].rows[0])
