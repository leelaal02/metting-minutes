"""커밋된 빈 양식 에셋(표·문단)의 자동 매핑 왕복 회귀 테스트.

make_example_forms.py가 만드는 두 양식은 [4] 자동 매핑 경로의 유일한 커밋 에셋이다.
여기서 inspect(구조 확인) → apply(토큰 삽입) → render(값 채움) 왕복이 실제 값으로
채워지는지 검증해, 문서로만 존재하던 자동 매핑 경로에 실행 가능한 안전망을 둔다.

매핑(mapping.json)은 원래 Claude가 만드는 비결정 산출물이라, 여기서는 각 양식의
알려진 구조에 맞춰 손으로 고정 매핑을 짜서 스크립트 부분만 검증한다.
"""
import json
from pathlib import Path

from docx import Document

from apply_form_mapping import apply_mapping
from inspect_template import inspect_template
from make_example_forms import build_paragraph_form, build_table_form
from render_docx_template import render_template

SAMPLE = Path(__file__).resolve().parent.parent / "examples" / "sample_minutes.json"
TEMPLATES = Path(__file__).resolve().parent.parent / "templates"
TABLE_FORM = TEMPLATES / "example-blank-table-form.docx"
PARA_FORM = TEMPLATES / "example-blank-paragraph-form.docx"


def _sample():
    return json.loads(SAMPLE.read_text(encoding="utf-8"))


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# --- 커밋 에셋이 존재하고 토큰이 없는지 -------------------------------------
def test_committed_forms_exist_and_are_tokenless():
    for form in (TABLE_FORM, PARA_FORM):
        assert form.exists(), f"빈 양식 에셋 누락: {form.name} (make_example_forms.py 재생성 필요)"
        assert inspect_template(str(form))["has_tokens"] is False


def test_generator_matches_committed_structure(tmp_path):
    """생성기 출력과 커밋 에셋의 구조(라벨/문단)가 일치한다(에셋이 낡지 않도록)."""
    gen_table = tmp_path / "t.docx"
    build_table_form().save(str(gen_table))
    assert [c["text"] for c in inspect_template(str(gen_table))["tables"][0]["cells"]] == [
        c["text"] for c in inspect_template(str(TABLE_FORM))["tables"][0]["cells"]
    ]
    gen_para = tmp_path / "p.docx"
    build_paragraph_form().save(str(gen_para))
    assert [p["text"] for p in inspect_template(str(gen_para))["paragraphs"]] == [
        p["text"] for p in inspect_template(str(PARA_FORM))["paragraphs"]
    ]


# --- 표 양식 왕복 -----------------------------------------------------------
def test_table_form_roundtrip_fills_values(tmp_path):
    tokenized = tmp_path / "tok.docx"
    final = tmp_path / "final.docx"
    mapping = {
        "table": 0,
        "fills": [
            {"row": 0, "col": 1, "mode": "inline", "fields": ["title"]},
            {"row": 1, "col": 1, "mode": "inline", "fields": ["date"]},
            {"row": 2, "col": 1, "mode": "inline", "fields": ["attendees"]},
            {"row": 3, "col": 1, "mode": "block", "fields": ["purpose"]},
            {"row": 4, "col": 1, "mode": "block", "fields": ["discussion"]},
            {"row": 5, "col": 1, "mode": "block", "fields": ["decisions"]},
            {"row": 6, "col": 1, "mode": "block", "fields": ["open_issues"]},
            {"row": 7, "col": 1, "mode": "block", "fields": ["action_items"]},
            {"row": 8, "col": 1, "mode": "block", "fields": ["next_meeting"]},
        ],
    }
    apply_mapping(str(TABLE_FORM), mapping, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    text = _all_text(Document(str(final)))
    assert "2026 3분기 제품 로드맵 회의" in text          # title
    assert "김수민, 이정우, 박서연" in text                # attendees_joined
    assert "docx 변환은 python-docx로 진행" in text        # decisions 반복
    assert "실시간 STT 도입 시점 미정" in text             # open_issues 반복
    assert "python-docx 렌더러 PoC 작성" in text           # action_items 반복


# --- 문단 양식 왕복 ---------------------------------------------------------
def test_paragraph_form_roundtrip_fills_values(tmp_path):
    tokenized = tmp_path / "tok.docx"
    final = tmp_path / "final.docx"
    mapping = {
        "paragraphs": [
            {"para": 2, "mode": "inline", "fields": ["title"]},
            {"para": 3, "mode": "inline", "fields": ["date"]},
            {"para": 4, "mode": "inline", "fields": ["attendees"]},
            {"para": 5, "mode": "inline", "fields": ["purpose"]},
            {"para": 7, "mode": "block", "fields": ["discussion"]},
            {"para": 9, "mode": "block", "fields": ["decisions", "open_issues", "action_items"]},
            {"para": 11, "mode": "block", "fields": ["next_meeting", "notes"]},
        ]
    }
    apply_mapping(str(PARA_FORM), mapping, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    text = _all_text(Document(str(final)))
    assert "2026 3분기 제품 로드맵 회의" in text          # title inline
    assert "docx 변환은 python-docx로 진행" in text        # decisions block
    assert "python-docx 렌더러 PoC 작성" in text           # action_items block
    # 복수 field block은 섹션 라벨로 구분된다
    assert "[결정 사항]" in text
    assert "[미결 사항]" in text
    assert "[실행 항목]" in text
    assert "STT 엔진 로컬 vs 클라우드 선택 보류" in text   # open_issues 반복
    # 양식의 헤딩은 보존된다
    assert "2. 회의 내용" in text
