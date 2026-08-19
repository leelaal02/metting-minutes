import json
from pathlib import Path

from docx import Document

from render_docx import render_docx

SAMPLE = Path(__file__).resolve().parent.parent / "examples" / "sample_minutes.json"


def _sample():
    return json.loads(SAMPLE.read_text(encoding="utf-8"))


def test_creates_valid_docx(tmp_path):
    out = tmp_path / "out.docx"
    render_docx(_sample(), str(out))
    assert out.exists()
    doc = Document(str(out))  # 열리면 유효한 docx
    texts = [p.text for p in doc.paragraphs]
    assert "2026 3분기 제품 로드맵 회의" in texts


def test_headings_present(tmp_path):
    out = tmp_path / "out.docx"
    render_docx(_sample(), str(out))
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    for heading in ["참석자", "회의 목적", "논의 내용", "결정 사항",
                    "실행 항목 Action Items", "다음 회의 일정", "기타·특이사항"]:
        assert heading in texts


def test_action_items_table_has_rows(tmp_path):
    out = tmp_path / "out.docx"
    render_docx(_sample(), str(out))
    doc = Document(str(out))
    assert len(doc.tables) == 1
    # 헤더 1행 + action_items 2행
    assert len(doc.tables[0].rows) == 3


def test_null_next_meeting(tmp_path):
    data = _sample()
    data["next_meeting"] = None
    out = tmp_path / "out.docx"
    render_docx(data, str(out))
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "(미정)" in texts
