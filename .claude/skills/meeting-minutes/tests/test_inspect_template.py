"""inspect_template.py 테스트: 좌표·is_empty·병합 원점 dedup·음영·넘버링·has_tokens."""
import pytest
from docx import Document
from docx.oxml.ns import qn

from inspect_template import inspect_template


def _shade(cell, fill="F3F3F3", theme_fill=None):
    """셀에 배경 음영(w:shd)을 준다 — 양식의 색칠된 라벨 칸 재현.

    `theme_fill`을 주면 테마 색으로 칠한 칸(w:fill 없음)을 재현한다.
    """
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    if theme_fill is not None:
        shd.set(qn("w:themeFill"), theme_fill)
    else:
        shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _number(paragraph, num_id=1):
    """문단에 자동 번호(w:numPr)를 건다 — 빈 "1." "2." 자리 재현."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn("w:numPr"), {})
    ilvl = pPr.makeelement(qn("w:ilvl"), {})
    ilvl.set(qn("w:val"), "0")
    numId = pPr.makeelement(qn("w:numId"), {})
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def _make_form(path, *, tokens=False):
    """토큰 없는(또는 있는) 표 서식 최소 양식을 만든다.

    구조:
      row0: 병합된 제목 셀 "회 의 록"
      row1: "제 목 :" | (빈 값 칸)
      row2: "일 시 :" | "참 가 자"
      row3: (빈칸, tokens=True면 토큰 삽입) | (빈칸)
    """
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "회 의 록"
    table.cell(0, 0).merge(table.cell(0, 1))  # 가로 병합
    table.cell(1, 0).text = "제 목 :"
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "일 시 :"
    table.cell(2, 1).text = "참 가 자"
    if tokens:
        table.cell(3, 0).text = "{{ title }}"
    doc.save(str(path))


def test_no_tokens_detected(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p, tokens=False)
    result = inspect_template(str(p))
    assert result["has_tokens"] is False


def test_tokens_detected(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p, tokens=True)
    result = inspect_template(str(p))
    assert result["has_tokens"] is True


def test_tokens_in_header_detected(tmp_path):
    """머리말에만 토큰이 있어도 탐지한다 — 본문만 훑으면 놓친다."""
    doc = Document()
    doc.add_table(rows=1, cols=1)  # 본문엔 토큰 없음
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{ title }}"
    p = tmp_path / "header_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_tokens_in_footer_detected(tmp_path):
    """꼬리말에만 토큰이 있어도 탐지한다."""
    doc = Document()
    doc.add_paragraph("본문 문단")
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "{% for a in attendees %}{{ a }}{% endfor %}"
    p = tmp_path / "footer_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_tokens_in_nested_table_detected(tmp_path):
    """표 셀 안의 중첩표에 든 토큰도 탐지한다 — cell.text는 중첩표를 안 훑는다."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "{{ purpose }}"
    p = tmp_path / "nested_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_table_dimensions(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    table = inspect_template(str(p))["tables"][0]
    assert table["index"] == 0
    assert table["rows"] == 4
    assert table["cols"] == 2


def test_merged_origin_deduped(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    cells = inspect_template(str(p))["tables"][0]["cells"]
    # row0은 가로 병합 → 원점 (0,0) 하나만, merged=True
    row0 = [c for c in cells if c["row"] == 0]
    assert len(row0) == 1
    assert row0[0]["col"] == 0
    assert row0[0]["merged"] is True
    assert row0[0]["text"] == "회 의 록"


def test_empty_cell_flagged(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    cells = inspect_template(str(p))["tables"][0]["cells"]
    by_coord = {(c["row"], c["col"]): c for c in cells}
    assert by_coord[(1, 1)]["is_empty"] is True   # 빈 값 칸
    assert by_coord[(1, 0)]["is_empty"] is False  # "제 목 :"
    assert by_coord[(2, 1)]["text"] == "참 가 자"


def test_non_docx_raises(tmp_path):
    bad = tmp_path / "form.txt"
    bad.write_text("not a docx", encoding="utf-8")
    with pytest.raises(ValueError, match="docx"):
        inspect_template(str(bad))


def test_no_tables_returns_empty(tmp_path):
    doc = Document()
    doc.add_paragraph("표 없는 문단전용 양식")
    p = tmp_path / "noform.docx"
    doc.save(str(p))
    result = inspect_template(str(p))
    assert result["tables"] == []
    assert result["has_tokens"] is False


def test_shaded_label_cell_flagged(tmp_path):
    """회색 라벨 칸은 shaded=True, 흰 값 칸은 False — 값 자리 판별의 핵심 신호."""
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "회 의 안 건"
    _shade(table.cell(0, 0))
    table.cell(1, 0).text = ""  # 값 자리(흰 칸)
    p = tmp_path / "shaded.docx"
    doc.save(str(p))
    cells = inspect_template(str(p))["tables"][0]["cells"]
    by_coord = {(c["row"], c["col"]): c for c in cells}
    assert by_coord[(0, 0)]["shaded"] is True
    assert by_coord[(1, 0)]["shaded"] is False


@pytest.mark.parametrize("fill", ["F3F3F3", "4472C4", "FFFF00", "C6E0B4", "000000"])
def test_any_fill_color_is_shaded(tmp_path, fill):
    """라벨 칸이 회색이 아니어도(파랑·노랑·연두·검정) 음영으로 잡는다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _shade(table.cell(0, 0), fill=fill)
    p = tmp_path / f"fill_{fill}.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["tables"][0]["cells"][0]["shaded"] is True


def test_theme_fill_is_shaded(tmp_path):
    """테마 색으로 칠한 칸은 w:fill이 없다 — fill만 보면 놓친다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _shade(table.cell(0, 0), theme_fill="accent1")
    p = tmp_path / "theme.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["tables"][0]["cells"][0]["shaded"] is True


def test_white_fill_is_not_shaded(tmp_path):
    """흰색(FFFFFF)·auto 채움은 음영으로 보지 않는다."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _shade(table.cell(0, 0), fill="FFFFFF")
    p = tmp_path / "white.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["tables"][0]["cells"][0]["shaded"] is False


def test_empty_numbered_paragraph_flagged(tmp_path):
    """글자 없는 번호 문단은 여백이 아니라 값 자리 — numbered=True로 구분한다."""
    doc = Document()
    doc.add_paragraph("")          # 간격용 빈 문단
    numbered = doc.add_paragraph("")  # "1." 만 찍히는 값 자리
    _number(numbered)
    p = tmp_path / "numbered.docx"
    doc.save(str(p))
    paras = {x["index"]: x for x in inspect_template(str(p))["paragraphs"]}
    assert paras[0]["is_empty"] is True and paras[0]["numbered"] is False
    assert paras[1]["is_empty"] is True and paras[1]["numbered"] is True


def test_blocks_report_table_paragraph_order(tmp_path):
    """표와 문단의 본문 순서를 낸다 — 표 밖 문단이 어느 표 뒤인지 알아야 한다."""
    doc = Document()
    doc.add_paragraph("회 의 록")
    doc.add_table(rows=1, cols=1)
    doc.add_paragraph("")
    doc.add_table(rows=1, cols=1)
    p = tmp_path / "order.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["blocks"] == [
        {"type": "paragraph", "index": 0},
        {"type": "table", "index": 0},
        {"type": "paragraph", "index": 1},
        {"type": "table", "index": 1},
    ]


def test_paragraphs_dumped_with_index(tmp_path):
    doc = Document()
    doc.add_paragraph("개요")
    doc.add_paragraph("ㅇ (목적)")
    doc.add_paragraph("")  # 빈 문단도 인덱스 정확성을 위해 포함
    doc.add_paragraph("ㅇ (일시)")
    p = tmp_path / "para_form.docx"
    doc.save(str(p))
    paras = inspect_template(str(p))["paragraphs"]
    by_idx = {x["index"]: x for x in paras}
    assert by_idx[0]["text"] == "개요"
    assert by_idx[1]["text"] == "ㅇ (목적)"
    assert by_idx[2]["is_empty"] is True
    assert by_idx[3]["text"] == "ㅇ (일시)"
